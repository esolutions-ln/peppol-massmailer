#!/usr/bin/env python3
"""Generate InvoiceDirect Developer Guide as a Word document."""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import datetime

doc = Document()

# ── Style setup ──────────────────────────────────────────────────────────────
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(10.5)
style.paragraph_format.space_after = Pt(4)
style.paragraph_format.line_spacing = 1.15

for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)
    hs.font.name = 'Calibri'

doc.styles['Heading 1'].font.size = Pt(18)
doc.styles['Heading 2'].font.size = Pt(14)
doc.styles['Heading 3'].font.size = Pt(12)

# Helper functions
def add_code_block(doc, code_text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(code_text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x2D, 0x2D, 0x2D)
    # light grey shading
    shading = p.paragraph_format.element.makeelement(qn('w:shd'), {
        qn('w:fill'): 'F2F2F2',
        qn('w:val'): 'clear',
    })
    p.paragraph_format.element.get_or_add_pPr().append(shading)

def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.left_indent = Cm(1.27 + level * 0.63)
    return p

def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(9)
                run.font.bold = True
    for row_data in rows:
        row = table.add_row()
        for i, val in enumerate(row_data):
            cell = row.cells[i]
            cell.text = str(val)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)
    return table

def add_note(doc, text, label="Note"):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(f"{label}: ")
    run.font.bold = True
    run.font.size = Pt(9.5)
    run = p.add_run(text)
    run.font.size = Pt(9.5)
    run.font.italic = True

def add_warning(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run("⚠ ")
    run.font.size = Pt(9.5)
    run = p.add_run(text)
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(0xCC, 0x66, 0x00)

def add_api_endpoint(doc, method, path, description, auth=None):
    p = doc.add_paragraph()
    run = p.add_run(f"{method} ")
    run.font.bold = True
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(0x00, 0x66, 0x99)
    run = p.add_run(path)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    p2 = doc.add_paragraph(description)
    p2.paragraph_format.left_indent = Cm(0.5)
    if auth:
        add_note(doc, auth, "Auth")

# ── Title Page ───────────────────────────────────────────────────────────────
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('InvoiceDirect Platform')
run.font.size = Pt(28)
run.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)
run.font.bold = True

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('Developer Integration Guide')
run.font.size = Pt(20)
run.font.color.rgb = RGBColor(0x4A, 0x6F, 0xA5)

doc.add_paragraph()
desc = doc.add_paragraph()
desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = desc.add_run('Integrate your ERP system to send fiscalised invoices\nvia Email (PDF) and PEPPOL e-Delivery (BIS 3.0 / AS4)')
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
run.font.italic = True

doc.add_paragraph()
doc.add_paragraph()
ver = doc.add_paragraph()
ver.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = ver.add_run(f'Document Version 1.0 — {datetime.date.today().strftime("%B %Y")}')
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

doc.add_page_break()

# ── Table of Contents placeholder ────────────────────────────────────────────
doc.add_heading('Table of Contents', level=1)
toc_items = [
    "1. Platform Overview",
    "2. Getting Started",
    "   2.1 Authentication & API Keys",
    "   2.2 Base URL & Environments",
    "   2.3 Common Request Patterns",
    "3. Emailer — PDF Invoice Delivery",
    "   3.1 Architecture Overview",
    "   3.2 Sending a Single Invoice",
    "   3.3 Mass Campaign Dispatch",
    "   3.4 ERP-Driven Dispatch (Dual-Channel Router)",
    "   3.5 PDF Sources & Validation",
    "   3.6 Email Templates",
    "   3.7 Campaign Completion Webhook",
    "   3.8 PDF Folder Watcher",
    "   3.9 Dashboard & Reporting",
    "4. PEPPOL e-Delivery — BIS 3.0 / AS4",
    "   4.1 PEPPOL 4-Corner Model",
    "   4.2 Organization Onboarding",
    "   4.3 Certificate Lifecycle",
    "   4.4 eRegistry — Access Points & Participant Links",
    "   4.5 Sending via PEPPOL",
    "   4.6 Receiving via PEPPOL (C3 / C4)",
    "   4.7 SMP Service Metadata Publishing",
    "   4.8 SML DNS Resolution",
    "   4.9 AS4 Sign-Then-Encrypt Transport",
    "   4.10 PEPPOL Invitation Flow",
    "5. ERP Adapters",
    "   5.1 Architecture (Hexagonal Ports & Adapters)",
    "   5.2 Sage Intacct",
    "   5.3 Sage Network",
    "   5.4 QuickBooks Online",
    "   5.5 Microsoft Dynamics 365",
    "   5.6 Odoo",
    "   5.7 Generic API",
    "6. Billing & Metering",
    "7. Webhooks & Callbacks",
    "8. Error Handling & Retry",
    "9. Security Overview",
    "10. Deployment"
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(1)
    p.runs[0].font.size = Pt(10)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 1 — Platform Overview
# ═══════════════════════════════════════════════════════════════════════════════
doc.add_heading('1. Platform Overview', level=1)

doc.add_paragraph(
    'InvoiceDirect is a multi-tenant, dual-channel invoice delivery platform. '
    'It enables organisations to send fiscalised invoices to their customers via '
    'two complementary channels:'
)

channels = [
    ('Email with PDF Attachment', 'Traditional email delivery with a PDF invoice attached. '
     'The PDF is produced externally (ERP system, fiscal device, or manual upload) and '
     'the platform handles template rendering, delivery, and tracking.'),
    ('PEPPOL e-Delivery (BIS 3.0 / AS4)', 'Electronic invoice exchange using the PEPPOL '
     '4-corner model. Invoices are transformed to UBL 2.1 XML (BIS Billing 3.0), '
     'validated against the EN 16931 Schematron, signed and encrypted via AS4, and '
     'delivered to the buyer\'s Access Point.'),
]
for title_text, desc_text in channels:
    p = doc.add_paragraph()
    run = p.add_run(f'{title_text}: ')
    run.font.bold = True
    run = p.add_run(desc_text)

doc.add_paragraph()
doc.add_heading('Key Capabilities', level=2)
caps = [
    'Mass campaign dispatch with Java 25 virtual-thread parallelism',
    'Dual-channel router — per-customer EMAIL or PEPPOL (BOTH supported)',
    'PEPPOL Access Point (AP) gateway — SMP publishing, SML DNS resolution, AS4 sign-then-encrypt',
    'ERP adapter layer — Sage Intacct, QuickBooks Online, Dynamics 365, Odoo, Generic API',
    'PDF inbox watcher — file-system monitoring for automated dispatch',
    'Fiscal validation — ZIMRA fiscal marker verification (fdms.zimra.co.zw)',
    'Billing & metering — tiered rate profiles with monthly auto-invoicing',
    'Campaign completion webhooks — HMAC-SHA256 signed callbacks',
    'OpenAPI 3.0 (Swagger) documentation — available per environment',
]
for c in caps:
    add_bullet(doc, c)

doc.add_heading('Technology Stack', level=2)
add_table(doc,
    ['Component', 'Technology'],
    [
        ['Language', 'Java 25 (virtual threads via Project Loom)'],
        ['Framework', 'Spring Boot 4.0 / Spring Framework 7.0'],
        ['Database', 'PostgreSQL 16+ (H2 for local development)'],
        ['ORM', 'Spring Data JPA / Hibernate 6'],
        ['Templating', 'Thymeleaf 3 (emails) + React 18 / TypeScript / Vite (frontend)'],
        ['PDF Validation', 'Magic-byte (25 50 44 46 2D) + ZIMRA fiscal marker scanning'],
        ['PEPPOL Schematron', 'ph-schematron-xslt 9.1.1 (ISO 3-phase pipeline, 160 assertions)'],
        ['AS4 Security', 'Apache Santuario 4.0 (XML-DSIG + XML-Enc, AES-256 / RSA-OAEP)'],
        ['SMTP', 'Spring JavaMailSender (fallback) + Brevo REST API (preferred)'],
        ['OpenAPI', 'Springdoc 3.0 with Swagger UI'],
        ['Deployment', 'Native JAR + systemd + nginx, or Docker Compose'],
    ])

doc.add_heading('Architecture Diagram', level=2)
doc.add_paragraph(
    '(Conceptual architecture — see individual sections for detailed flow diagrams)'
)
add_code_block(doc, """\
┌─────────────┐    ┌──────────────────────────────────────────────┐    ┌─────────────┐
│   Your ERP   │───▶│           InvoiceDirect Platform            │───▶│  Customers  │
│ (Sage/QB/    │    │                                              │    │             │
│  D365/Odoo)  │    │  ┌─────────┐  ┌──────────┐  ┌───────────┐  │    │  ┌────────┐ │
└─────────────┘    │  │  ERP     │──▶  Campaign  │  │  EMAIL    │───────│  Email  │ │
                   │  │ Adapters │  │ Orchestr. │  │  (PDF)    │  │    │  Inbox  │ │
                   │  └─────────┘  └──────────┘  └───────────┘  │    │  └────────┘ │
                   │                              ┌───────────┐  │    │             │
                   │                              │  PEPPOL   │───────│  ERP       │
                   │                              │  (BIS 3.0 │  │    │  (C4)      │
                   │                              │   / AS4)  │  │    └────────────┘
                   │                              └───────────┘  │
                   └──────────────────────────────────────────────┘""")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 2 — Getting Started
# ═══════════════════════════════════════════════════════════════════════════════
doc.add_heading('2. Getting Started', level=1)

doc.add_heading('2.1 Authentication & API Keys', level=2)
doc.add_paragraph(
    'All API requests (except public PEPPOL endpoints) require authentication. '
    'The platform supports two auth modes:'
)
add_bullet(doc, 'API Key: Pass via the X-API-Key header. Generate or rotate keys via the dashboard or POST /api/v1/my/rotate-api-key.')
add_bullet(doc, 'Session Token: Obtain via POST /api/v1/admin/login (admin) or POST /api/v1/org/login (org member).')
doc.add_paragraph(
    'PEPPOL AS4 inbound (/peppol/as4/receive) uses HMAC-SHA256 with a pre-shared secret. '
    'Public endpoints (/bdxr/smp/**, /peppol/as4/health) require no authentication.'
)

doc.add_heading('2.2 Base URL & Environments', level=2)
add_table(doc,
    ['Environment', 'Base URL', 'Purpose'],
    [
        ['Production', 'https://ap.invoicedirect.biz', 'Live invoice delivery'],
        ['Staging / Test', 'https://staging.invoicedirect.biz', 'Integration testing'],
        ['Local Development', 'http://localhost:9199', 'Local Docker Compose'],
    ])

doc.add_heading('Common Headers', level=2)
add_table(doc,
    ['Header', 'Value', 'Required'],
    [
        ['Content-Type', 'application/json', 'Yes (all API requests)'],
        ['X-API-Key', '{your-api-key}', 'Yes (org-scoped)'],
        ['Authorization', 'Bearer {session-token}', 'Alternative (session-based)'],
        ['X-PEPPOL-Sender-ID', '9915:org-slug', 'Required for AS4 inbound'],
        ['X-PEPPOL-Receiver-ID', '0190:ZW123456789', 'Required for AS4 inbound'],
        ['X-PEPPOL-Signature', 'HMAC-SHA256 hex digest', 'Required for AS4 inbound'],
    ])

doc.add_heading('2.3 Common Request Patterns', level=2)
doc.add_paragraph(
    'All API endpoints live under /api/v1/. Responses follow a consistent JSON envelope. '
    'Error responses include a message field and optional details array.'
)
add_code_block(doc, """\
// Success
HTTP 200 / 201
{ "id": "uuid", "status": "delivered", "messageId": "<msg-id>", ... }

// Error
HTTP 400 / 404 / 502
{ "message": "Description of what went wrong", "details": [...] }""")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 3 — Emailer (PDF Invoice Delivery)
# ═══════════════════════════════════════════════════════════════════════════════
doc.add_heading('3. Emailer — PDF Invoice Delivery', level=1)

doc.add_heading('3.1 Architecture Overview', level=2)
doc.add_paragraph(
    'The Emailer subsystem delivers fiscalised invoice PDFs to customer email inboxes. '
    'It supports synchronous single-send, asynchronous mass campaigns, and fully automated '
    'ERP-driven dispatch. The platform does NOT generate PDFs — they are produced by the '
    'originating ERP system, fiscal device, or external invoice printing system.'
)

add_code_block(doc, """\
                PDF Sources                          Email Channels
┌──────────────┐    ┌──────────────┐         ┌──────────────────┐
│  ERP System  │───▶│  File Path   │────────▶│  Brevo REST API  │
│ (Sage/QB/    │    │  Base64 JSON │         │  (preferred)     │
│  D365/Odoo)  │    │  Upload      │         └──────────────────┘
└──────────────┘    └──────────────┘                 │
                                           ┌──────────────────┐
┌──────────────┐    ┌──────────────┐         │  JavaMail SMTP   │
│  PDF Watcher │───▶│  Folder      │────────▶│  (fallback)      │
│  Filesystem  │    │  Monitor     │         └──────────────────┘
└──────────────┘    └──────────────┘""")

doc.add_paragraph(
    'Flow: (1) PDF arrives via file path, Base64 in JSON payload, multipart upload, or filesystem watcher. '
    '(2) PDF is validated (magic bytes, size cap, optional ZIMRA fiscal markers). '
    '(3) Email body is rendered (Thymeleaf template or user-defined {{placeholder}} template). '
    '(4) Email is sent via Brevo HTTPS API or SMTP fallback. '
    '(5) Delivery is recorded for metering and billing.'
)

doc.add_heading('3.2 Sending a Single Invoice', level=2)
doc.add_paragraph(
    'Use the synchronous endpoint for immediate single-invoice delivery. '
    'The PDF attachment can be provided as a Base64 string, file path, or multipart file upload.'
)

doc.add_heading('Single Invoice (Base64 or File Path)', level=3)
add_api_endpoint(doc, 'POST', '/api/v1/mail/invoice', 'Send a single fiscalised invoice with PDF attachment.')
add_code_block(doc, """\
{
  "to": "buyer@acmecorp.co.zw",
  "recipientName": "Acme Corporation",
  "subject": "Invoice INV-2024-001 from eSolutions",
  "templateName": "invoice",
  "invoiceNumber": "INV-2024-001",
  "invoiceDate": "2024-12-01",
  "dueDate": "2024-12-31",
  "totalAmount": 2300.00,
  "vatAmount": 300.00,
  "currency": "USD",
  "pdfBase64": "JVBERi0xLjQKMSAwIG9iago8PAovVHlwZSAvQ2F0YWxvZw...",
  "pdfFileName": "INV-2024-001.pdf",
  "variables": {
    "company": "eSolutions (Pvt) Ltd",
    "customerAccount": "CUST-001"
  }
}""")

doc.add_heading('Single Invoice (Multipart Upload)', level=3)
add_api_endpoint(doc, 'POST', '/api/v1/mail/invoice/upload',
    'Same as above but PDF sent as multipart file upload.')
doc.add_paragraph(
    'Form fields: metadata (same JSON as above, without pdfBase64), file (the PDF).'
)

doc.add_heading('Response Format', level=3)
add_code_block(doc, """\
HTTP 200 (delivered)
{ "status": "delivered", "recipient": "buyer@acmecorp.co.zw",
  "invoiceNumber": "INV-2024-001", "messageId": "<20241201-example@mail>",
  "retryable": false }

HTTP 400 (validation error)
{ "status": "failed", "recipient": "buyer@acmecorp.co.zw",
  "invoiceNumber": "INV-2024-001", "error": "PDF exceeds maximum size (10 MB)",
  "retryable": false }

HTTP 502 (send failure)
{ "status": "failed", "recipient": "buyer@acmecorp.co.zw",
  "invoiceNumber": "INV-2024-001", "error": "Could not connect to SMTP server",
  "retryable": true }""")

doc.add_heading('3.3 Mass Campaign Dispatch', level=2)
doc.add_paragraph(
    'For sending invoices in bulk, create a campaign. '
    'Dispatch is fully asynchronous — the endpoint returns immediately with the campaign ID, '
    'and recipients are processed in parallel batches using Java 25 virtual threads.'
)

add_api_endpoint(doc, 'POST', '/api/v1/campaigns',
    'Create and queue a mass invoice campaign.')
add_code_block(doc, """\
{
  "organizationId": "org-uuid",
  "name": "December 2024 Invoices",
  "subject": "Your Invoice from eSolutions",
  "templateName": "invoice",
  "callbackUrl": "https://erp.acmecorp.co.zw/webhooks/campaign-completed",
  "recipients": [
    {
      "email": "buyer1@acmecorp.co.zw",
      "recipientName": "Acme Corp",
      "invoiceNumber": "INV-2024-001",
      "totalAmount": 2300.00,
      "vatAmount": 300.00,
      "currency": "USD",
      "pdfBase64": "JVBERi0xLjQ...",
      "pdfFileName": "INV-2024-001.pdf",
      "mergeFields": { "purchaseOrder": "PO-12345" }
    },
    {
      "email": "buyer2@example.co.zw",
      "recipientName": "Example Ltd",
      "invoiceNumber": "INV-2024-002",
      "totalAmount": 1150.00,
      "vatAmount": 150.00,
      "currency": "USD",
      "pdfFilePath": "/var/lib/invoices/INV-2024-002.pdf"
    }
  ]
}""")

doc.add_heading('Campaign Status Polling', level=3)
add_api_endpoint(doc, 'GET', '/api/v1/campaigns/{id}',
    'Poll campaign progress and final status.')
add_code_block(doc, """\
HTTP 200
{
  "id": "campaign-uuid",
  "name": "December 2024 Invoices",
  "status": "COMPLETED",
  "totalRecipients": 50,
  "sentCount": 48,
  "failedCount": 1,
  "skippedCount": 1,
  "createdAt": "2024-12-01T10:00:00Z",
  "completedAt": "2024-12-01T10:05:30Z"
}""")

doc.add_heading('Retry Failed Recipients', level=3)
add_api_endpoint(doc, 'POST', '/api/v1/campaigns/{id}/retry',
    'Retry all failed recipients in a campaign.')

doc.add_heading('Campaign List', level=3)
add_api_endpoint(doc, 'GET', '/api/v1/campaigns',
    'List all campaigns (filters: status, date range).')

doc.add_heading('3.4 ERP-Driven Dispatch (Dual-Channel Router)', level=2)
doc.add_paragraph(
    'This is the primary integration point for ERP systems. A single dispatch request can '
    'contain invoices destined for both EMAIL and PEPPOL recipients — the platform '
    'automatically routes each invoice to the correct channel based on the customer\'s '
    'delivery mode (resolved via the customer\'s PEPPOL participant link or org default).'
)

doc.add_paragraph(
    'The dispatch flow: (1) Fetch invoice metadata from the ERP system (or provide inline in the request). '
    '(2) For each invoice, resolve the customer\'s delivery mode. '
    '(3) EMAIL customers: attach PDF and send via email. '
    '(4) PEPPOL customers: build UBL 2.1 XML, validate Schematron, transmit via AS4.'
)

add_api_endpoint(doc, 'POST', '/api/v1/erp/dispatch',
    'Fetch invoices from ERP and dispatch via dual-channel router.')
add_code_block(doc, """\
{
  "organizationId": "org-uuid",
  "erpSource": "SAGE_INTACCT",
  "invoiceIds": ["INV-001", "INV-002", "INV-003"],
  "campaignName": "Dec 2024 Auto-Dispatch"
}""")

doc.add_heading('3.5 PDF Sources & Validation', level=2)

doc.add_heading('PDF Attachment Sources', level=3)
add_table(doc,
    ['Source', 'Field', 'Description'],
    [
        ['Base64', 'pdfBase64', 'Base64-encoded PDF bytes in the JSON payload. Max 10 MB decoded.'],
        ['File Path', 'pdfFilePath', 'Absolute path to PDF on shared filesystem. Must be under pdf-inbox-base-path.'],
        ['Multipart Upload', 'file', 'PDF file uploaded as multipart/form-data.'],
        ['PDF Watcher', 'N/A', 'Filesystem monitoring — see Section 3.8.'],
        ['ERP Adapter', 'N/A', 'Adapter fetches PDF from ERP\'s API or export directory.'],
    ])

doc.add_heading('PDF Validation Pipeline', level=3)
doc.add_paragraph(
    'Every PDF undergoes two validation stages before dispatch:'
)
add_bullet(doc, 'Structural Validation: Magic-byte check (%PDF- header = 25 50 44 46 2D) and size cap (default 10 MB, configurable via MAX_PDF_BYTES).')
add_bullet(doc, 'Fiscal Validation (optional, enabled by default): Scans PDF content for ZIMRA fiscal markers — must contain fdms.zimra.co.zw domain, 16-char hex verification code, and at least one of (Device ID, Fiscal Day, Fiscal Invoice Number, Global Receipt Number).')

add_code_block(doc, """\
// Fiscal marker extraction from PDF
"ZIMRA Verification: FD=20241201, DEV=T24A00001, GIC=42,
 VC=A1B2C3D4E5F6G7H8, QR=https://fdms.zimra.co.zw/verify/...\"""")

doc.add_heading('3.6 Email Templates', level=2)
doc.add_paragraph(
    'The platform supports two template systems:'
)

add_bullet(doc, 'System Templates: Thymeleaf HTML files under src/main/resources/templates/email/. Built-in templates: invoice.html (with fiscal verification box and QR code), platform-invoice.html (billing), generic.html, welcome.html, peppol-invitation.html.')
add_bullet(doc, 'Custom DB Templates: Organisation-scoped templates stored in the email_templates table. Use {{placeholder}} syntax for dynamic fields. CRUD via /api/v1/my/email-templates/.')

add_api_endpoint(doc, 'POST', '/api/v1/my/email-templates', 'Create a custom email template.')
add_code_block(doc, """\
{
  "name": "My Invoice Template",
  "subject": "Invoice {{invoiceNumber}} from {{company}}",
  "body": "<p>Dear {{recipientName}},</p><p>Your invoice is attached.</p>",
  "isDefault": true
}""")

doc.add_heading('3.7 Campaign Completion Webhook', level=2)
doc.add_paragraph(
    'When a campaign completes and a callbackUrl was provided, the platform sends an '
    'HMAC-SHA256 signed POST to that URL. This allows your ERP to be notified when '
    'all invoices in a batch have been processed.'
)

doc.add_paragraph('Webhook Payload:')
add_code_block(doc, """\
POST {campaign.callbackUrl}
Content-Type: application/json
X-Webhook-Signature: {HMAC-SHA256 hex digest}
X-Webhook-Event: campaign.completed

{
  "campaignId": "uuid",
  "name": "December 2024 Invoices",
  "status": "COMPLETED",
  "totalRecipients": 50,
  "sent": 48,
  "failed": 1,
  "skipped": 1,
  "completedAt": "2024-12-01T10:05:30Z"
}""")

doc.add_paragraph(
    'Verify the signature using the shared WEBHOOK_SECRET. '
    'Algorithm: HMAC-SHA256 of the raw JSON body.'
)

doc.add_heading('3.8 PDF Folder Watcher', level=2)
doc.add_paragraph(
    'The optional filesystem watcher monitors a directory for new PDF files and automatically '
    'dispatches them. Each PDF must have a companion .json sidecar file with metadata.'
)

doc.add_paragraph('Sidecar file format (same-name JSON next to the PDF):')
add_code_block(doc, """\
{
  "organizationId": "org-uuid",
  "invoiceNumber": "INV-2024-001",
  "recipientEmail": "buyer@acmecorp.co.zw",
  "recipientName": "Acme Corp",
  "totalAmount": 2300.00,
  "vatAmount": 300.00,
  "currency": "USD",
  "fiscalDeviceSerialNumber": "T24A00001",
  "globalInvoiceCounter": "42",
  "verificationCode": "A1B2C3D4E5F6G7H8",
  "erpSource": "SAGE_INTACCT"
}""")

doc.add_paragraph(
    'Enable via PDF_WATCHER_ENABLED=true and configure the inbox directory. '
    'Successfully dispatched PDFs move to a processed/ subfolder, failures to failed/.'
)

doc.add_heading('3.9 Dashboard & Reporting', level=2)
doc.add_paragraph(
    'Organisations have access to a full dashboard for monitoring invoice delivery, '
    'viewing campaign history, and tracking billing.'
)

add_api_endpoint(doc, 'GET', '/api/v1/my/stats',
    'Dashboard statistics (campaigns, deliveries, billing period cost).')
add_api_endpoint(doc, 'GET', '/api/v1/my/campaigns',
    'List campaigns with filters by status.')
add_api_endpoint(doc, 'GET', '/api/v1/my/campaigns/{id}',
    'Single campaign detail with all invoice records.')
add_api_endpoint(doc, 'GET', '/api/v1/my/invoices',
    'Paginated invoice list (filters: status, date range, invoice number).')
add_api_endpoint(doc, 'GET', '/api/v1/my/invoices/{invoiceNumber}',
    'Single invoice lookup by invoice number.')
add_api_endpoint(doc, 'GET', '/api/v1/my/customers/{customerId}/invoices',
    'All invoices for a specific customer.')
add_api_endpoint(doc, 'GET', '/api/v1/my/billing',
    'Current billing summary.')
add_api_endpoint(doc, 'GET', '/api/v1/my/billing/history',
    'Full billing history.')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 4 — PEPPOL e-Delivery
# ═══════════════════════════════════════════════════════════════════════════════
doc.add_heading('4. PEPPOL e-Delivery — BIS 3.0 / AS4', level=1)

doc.add_heading('4.1 PEPPOL 4-Corner Model', level=2)
doc.add_paragraph(
    'InvoiceDirect implements a full PEPPOL Access Point (AP) gateway operating in the '
    '4-corner model. It acts as both the sending AP (C2) for outbound invoices and the '
    'receiving AP (C3) for inbound invoices.'
)

add_code_block(doc, """\
                        PEPPOL 4-Corner Model
  Supplier (C1)                     Buyer's ERP (C4)
       │                                  ▲
       │  1. UBL Invoice                  │  6. UBL Invoice via
       │     (via API)                    │     C4 Webhook
       ▼                                  │
  ┌─────────────────────────────────────────────────┐
  │         InvoiceDirect AP Gateway (C2 + C3)       │
  │                                                  │
  │  2. Resolve Buyer AP via SML DNS/SMP             │
  │     └─▶ b-0190-zw987654321.sml.peppoltest.org    │
  │                                                  │
  │  3. Get Endpoint + Certificate from SMP          │
  │     └─▶ https://buyer-ap.com/peppol/as4/receive  │
  │                                                  │
  │  4. Build UBL + Validate Schematron              │
  │  5. Sign & Encrypt (AS4) / HTTP POST             │
  │     └─▶ https://buyer-ap.com/peppol/as4/receive  │
  │                                                  │
  │  === Inbound Path ===                            │
  │  6. Receive AS4 / HTTP POST                      │
  │  7. Verify HMAC + Dedup (SHA-256)                │
  │  8. Route to C4 (org's webhook)                  │
  └─────────────────────────────────────────────────┘""")

doc.add_heading('4.2 Organization Onboarding', level=2)
doc.add_paragraph(
    'Before sending or receiving PEPPOL invoices, an organisation must be onboarded. '
    'This is an admin operation that registers the org\'s participant ID and creates '
    'the Access Point entries.'
)

add_api_endpoint(doc, 'POST', '/api/v1/admin/orgs/{orgId}/peppol/onboard',
    'Onboard an organisation for PEPPOL delivery.')
add_code_block(doc, """\
// Request
{ "deliveryMode": "BOTH" }
// or deliveryMode: "AS4" for PEPPOL-only

// Response
{
  "organizationId": "org-uuid",
  "participantId": "0190:ZW12345678",
  "deliveryMode": "BOTH",
  "gatewayAccessPoint": {
    "id": "ap-uuid",
    "participantId": "9915:org-slug",
    "endpointUrl": "https://ap.invoicedirect.biz/peppol/as4/receive",
    "simplifiedHttpDelivery": true
  }
}""")

doc.add_paragraph(
    'The participant ID is derived from the organisation\'s VAT number (format 0190:ZW{VAT}). '
    'The sending participant ID uses the 9915 scheme with the org slug. '
    'Delivery modes: EMAIL (PDF only), AS4 (PEPPOL only), BOTH (dual-channel).'
)

doc.add_heading('4.3 Certificate Lifecycle', level=2)
doc.add_paragraph(
    'Each PEPPOL-enabled organisation requires an X.509 certificate for AS4 message signing '
    'and encryption. Certificates are managed via the admin PKI API.'
)

add_api_endpoint(doc, 'POST', '/api/v1/admin/orgs/{orgId}/peppol/certs',
    'Upload a certificate and private key.')
add_api_endpoint(doc, 'POST', '/api/v1/admin/orgs/{orgId}/peppol/certs/rotate',
    'Rotate the certificate — expires the current, stores the new.')
add_api_endpoint(doc, 'GET', '/api/v1/admin/orgs/{orgId}/peppol/certs/active',
    'Get the current active certificate (public key only).')
add_api_endpoint(doc, 'GET', '/api/v1/admin/orgs/{orgId}/peppol/certs',
    'List all certificates for the org.')

add_code_block(doc, """\
// Upload (the JSON must use embedded \\n for PEM newlines)

{
  "certificate": "-----BEGIN CERTIFICATE-----\\nMIIBxT...\\n-----END CERTIFICATE-----",
  "privateKey": "-----BEGIN PRIVATE KEY-----\\nMIIEvQ...\\n-----END PRIVATE KEY-----",
  "description": "PEPPOL AP Certificate 2024"
}

// Response (cert without private key)
{
  "id": "cert-uuid",
  "subjectDn": "CN=Test PEPPOL AP, O=MyOrg, C=ZW",
  "issuerDn": "CN=Test PEPPOL AP, O=MyOrg, C=ZW",
  "validFrom": "2024-12-01T00:00:00Z",
  "validTo": "2034-11-30T23:59:59Z",
  "status": "ACTIVE"
}""")

add_warning(doc, 'Private keys are not returned from any API endpoint. Store them securely after upload.')
doc.add_paragraph()
add_note(doc, 'PEM newlines must be sent as literal \\n escape sequences in JSON, not raw newlines. This is required by Jackson 3\'s strict character validation.')

doc.add_heading('4.4 eRegistry — Access Points & Participant Links', level=2)
doc.add_paragraph(
    'The eRegistry is the local database of PEPPOL Access Points and participant routing '
    'information. It serves the role of the PEPPOL SMP for participants hosted on this AP.'
)

doc.add_heading('Access Points', level=3)
add_api_endpoint(doc, 'POST', '/api/v1/eregistry/access-points',
    'Register an Access Point.')
add_api_endpoint(doc, 'GET', '/api/v1/eregistry/access-points',
    'List Access Points (filter by organizationId or role).')
add_api_endpoint(doc, 'GET', '/api/v1/eregistry/access-points/by-participant/{participantId}',
    'Get AP by participant ID.')
add_api_endpoint(doc, 'PATCH', '/api/v1/eregistry/access-points/{id}/status?status=SUSPENDED',
    'Suspend or reactivate an Access Point.')

add_code_block(doc, """\
// Register an Access Point
{
  "organizationId": "org-uuid",          // null for external APs
  "participantId": "0190:ZW123456789",   // format: {scheme}:{value}
  "participantName": "Acme Corp AP",
  "role": "RECEIVER",                    // SENDER, RECEIVER, or GATEWAY
  "endpointUrl": "https://ap.invoicedirect.biz/peppol/as4/receive",
  "simplifiedHttpDelivery": true,
  "deliveryAuthToken": null,
  "certificate": "-----BEGIN CERTIFICATE-----\\n..."
}""")

doc.add_heading('Participant Links', level=3)
doc.add_paragraph(
    'A participant link connects a customer contact to their PEPPOL participant ID and '
    'receiver Access Point. This is how the platform knows where to deliver PEPPOL '
    'invoices for each customer.'
)

add_api_endpoint(doc, 'POST', '/api/v1/eregistry/participant-links',
    'Link a customer contact to their PEPPOL Access Point.')
add_api_endpoint(doc, 'GET', '/api/v1/eregistry/participant-links?organizationId=...',
    'List participant links.')
add_api_endpoint(doc, 'DELETE', '/api/v1/eregistry/participant-links/{id}',
    'Remove a participant link.')

add_code_block(doc, """\
// Create participant link
{
  "organizationId": "org-uuid",
  "customerEmail": "buyer@acmecorp.co.zw",
  "participantId": "0190:ZW987654321",
  "receiverAccessPointId": "ap-uuid",
  "preferredChannel": "PEPPOL"
}

// The routing resolution:
// 1. Check customer contact's deliveryMode override
// 2. Fall back to organisation default deliveryMode
// 3. Fall back to EMAIL""")

doc.add_heading('4.5 Sending via PEPPOL', level=2)
doc.add_paragraph(
    'PEPPOL dispatch is triggered when the dual-channel router resolves a customer\'s '
    'delivery mode to AS4 or BOTH. The dispatch flow:'
)

add_bullet(doc, '1. Resolve sender (C1) and receiver (C4) organisation and participant IDs')
add_bullet(doc, '2. Resolve the receiver\'s Access Point via the participant link')
add_bullet(doc, '3. Build UBL 2.1 XML using UblInvoiceBuilder (PEPPOL BIS Billing 3.0)')
add_bullet(doc, '4. Validate against PEPPOL-EN16931-UBL.sch (160 Schematron assertions)')
add_bullet(doc, '5. Create delivery record with status TRANSMITTING')
add_bullet(doc, '6. Transmit: simplified HTTP POST or full AS4 sign-then-encrypt')
add_bullet(doc, '7. Record delivery result (DELIVERED / FAILED)')

doc.add_heading('Simplified HTTP Delivery', level=3)
doc.add_paragraph(
    'When the receiver AP supports simplified HTTP delivery, the platform sends the '
    'raw UBL XML as an HTTP POST. This is suitable for private networks or when '
    'both APs trust the transport layer.'
)

add_code_block(doc, """\
POST {receiver-endpoint-url}
Content-Type: application/xml; charset=utf-8
X-PEPPOL-Document-Type: urn:oasis:names:specification:ubl:schema:xsd:Invoice-2::Invoice##
  urn:cen.eu:en16931:2017#compliant#urn:fdc:peppol.eu:2017:poacc:billing:3.0::2.1
X-PEPPOL-Process: urn:fdc:peppol.eu:2017:poacc:billing:01:1.0
X-Invoice-Number: INV-2024-001
Authorization: Bearer {delivery-auth-token}  // if configured

{raw UBL XML body}""")

doc.add_heading('4.6 Receiving via PEPPOL (C3 / C4)', level=2)
doc.add_paragraph(
    'InvoiceDirect acts as a receiving Access Point (C3), accepting inbound PEPPOL '
    'messages from other APs and routing them to the receiving organisation\'s ERP (C4).'
)

doc.add_heading('C3 — Inbound Endpoint', level=3)
add_api_endpoint(doc, 'POST', '/peppol/as4/receive',
    'Receive inbound PEPPOL AS4 or simplified HTTP messages.',
    'Public endpoint — HMAC-SHA256 verified. Requires X-PEPPOL-Sender-ID, '
    'X-PEPPOL-Receiver-ID, and X-PEPPOL-Signature headers.')

doc.add_paragraph('Inbound processing steps:')
add_bullet(doc, '1. Shape validation — message must contain <Invoice or <CreditNote')
add_bullet(doc, '2. Sender AP lookup — reject if sender participant ID is unknown or suspended')
add_bullet(doc, '3. HMAC verification — constant-time comparison using pre-shared inboundSharedSecret')
add_bullet(doc, '4. SHA-256 deduplication — returns existing receipt on duplicate payload')
add_bullet(doc, '5. Receiver organisation resolution — via AccessPoint.organizationId')
add_bullet(doc, '6. Persist InboundDocument with status RECEIVED')

doc.add_heading('C4 — ERP Webhook Forwarding', level=3)
doc.add_paragraph(
    'A scheduled job (every 30 seconds) picks up RECEIVED documents and forwards them '
    'to the receiving organisation\'s configured C4 webhook URL. The raw UBL XML is '
    'POSTed to the URL stored in the organisation\'s c4WebhookUrl field.'
)

add_code_block(doc, """\
POST {org.c4WebhookUrl}
Content-Type: application/xml; charset=utf-8
Authorization: Bearer {org.c4WebhookAuthToken}  // if configured

{raw UBL XML of received Invoice/CreditNote}""")

doc.add_paragraph(
    'The C4 routing job retries up to 3 times with a 30-second interval between attempts. '
    'After exhausting retries, the document is marked ROUTING_FAILED. '
    'If no C4 webhook URL is configured, the document is skipped (not retried).'
)

doc.add_heading('Inbound Document Audit', level=3)
add_api_endpoint(doc, 'GET', '/peppol/as4/inbox',
    'Admin-only paginated view of all received documents.')
add_api_endpoint(doc, 'GET', '/peppol/as4/health',
    'AS4 endpoint health check — returns UP status. Public.')

doc.add_heading('4.7 SMP Service Metadata Publishing', level=2)
doc.add_paragraph(
    'The platform implements BDXR SMP v2 (OASIS Standard), publishing service metadata '
    'for every registered participant. Other Access Points query this endpoint to discover '
    'our endpoint URL and AS4 certificate for a given participant.'
)

add_api_endpoint(doc, 'GET', '/bdxr/smp/{participantId}/services/{documentTypeId}',
    'Get SMP service metadata for a participant.',
    'Public — no authentication required.')
add_api_endpoint(doc, 'GET', '/bdxr/smp/{scheme}/{value}/services/{documentTypeId}',
    'Same as above with split participant ID.')

add_code_block(doc, """\
// Example: GET /bdxr/smp/0190:ZW987654321/services/...
// Response (BDXR SMP v2 XML)

<?xml version="1.0" encoding="UTF-8"?>
<ns2:SignedServiceMetadata
    xmlns:ns2="http://docs.oasis-open.org/bdxr/ns/SMP/2016/05">
  <ns2:ServiceMetadata>
    <ns2:ServiceInformation>
      <ns2:ParticipantIdentifier scheme="iso6523-actorid-upis">
        0190:ZW987654321</ns2:ParticipantIdentifier>
      <ns2:DocumentIdentifier scheme="cenbii-procid-ubl">
        urn:oasis:names:specification:ubl:schema:xsd:Invoice-2::Invoice##
        urn:cen.eu:en16931:2017#compliant#urn:fdc:peppol.eu:2017:
        poacc:billing:3.0::2.1</ns2:DocumentIdentifier>
      <ns2:ProcessList>
        <ns2:Process>
          <ns2:ProcessIdentifier scheme="cenbii-procid-ubl">
            urn:fdc:peppol.eu:2017:poacc:billing:01:1.0
          </ns2:ProcessIdentifier>
          <ns2:ServiceEndpointList>
            <ns2:Endpoint transportProfile="peppol-transport-as4-v2_0">
              <ns2:EndpointURI>
                https://ap.invoicedirect.biz/peppol/as4/receive
              </ns2:EndpointURI>
              <ns2:Certificate>Base64-encoded X.509...</ns2:Certificate>
              <ns2:ServiceDescription>InvoiceDirect Gateway AP</ns2:ServiceDescription>
            </ns2:Endpoint>
          </ns2:ServiceEndpointList>
        </ns2:Process>
      </ns2:ProcessList>
    </ns2:ServiceInformation>
  </ns2:ServiceMetadata>
</ns2:SignedServiceMetadata>""")

doc.add_heading('4.8 SML DNS Resolution', level=2)
doc.add_paragraph(
    'When sending to a PEPPOL participant, the platform first attempts SML DNS resolution '
    'to find the receiver\'s SMP host:'
)

add_code_block(doc, """\
DNS Pattern: {dnsPrefix}-{scheme}-{value}.{domain}
Example:     b-0190-zw987654321.sml.peppoltest.org
                      │   │              │
                      │   │              └── SML domain (configurable)
                      │   └────────────────── Participant value
                      └────────────────────── DNS prefix (default "b")

Resolution Flow:
1. Build hostname: b-0190-zw987654321.sml.peppoltest.org
2. DNS lookup via InetAddress.getAllByName()
3. If resolved → query http://{ip}/bdxr/smp/{participant}/services/{docType}
4. If DNS fails → fall back to configured PEPPOL_SMP_BASE_URL""")

doc.add_paragraph(
    'Configuration: PEPPOL_SML_DOMAIN (default sml.peppoltest.org for test network, '
    'sml.peppol.net for production), PEPPOL_SML_DNS_PREFIX (default b).'
)

doc.add_heading('4.9 AS4 Sign-Then-Encrypt Transport', level=2)
doc.add_paragraph(
    'Full AS4 ebMS 3.0 with WS-Security is implemented for secure message exchange. '
    'The transport applies a sign-then-encrypt pattern in five phases:'
)

add_bullet(doc, '1. SOAP 1.2 Envelope — Build ebMS 3.0 UserMessage with MessageInfo, PartyInfo, CollaborationInfo, and PayloadInfo referencing the body')
add_bullet(doc, '2. XML-DSIG Signing — RSA-SHA256 enveloped signature over the SOAP body (#body). Sender\'s X.509 certificate embedded in KeyInfo. Inclusive canonicalization (C14N).')
add_bullet(doc, '3. XML-Enc Encryption — Ephemeral AES-256 key encrypts the SOAP body. Key is RSA-OAEP wrapped using receiver\'s public key. Manual wrapping for maximum compatibility.')
add_bullet(doc, '4. HTTP POST — Content-Type: application/soap+xml; charset=UTF-8')
add_bullet(doc, '5. MDN Parsing — Parse ebMS receipt (success) or error (failure) from response')

add_code_block(doc, """\
SOAP Envelope Structure (simplified):
<S12:Envelope>
  <S12:Header>
    <eb:Messaging>
      <eb:UserMessage>
        <eb:MessageInfo>...</eb:MessageInfo>
        <eb:PartyInfo>...</eb:PartyInfo>
        <eb:CollaborationInfo>...</eb:CollaborationInfo>
        <eb:PayloadInfo><eb:PartInfo href="#body"/></eb:PayloadInfo>
      </eb:UserMessage>
    </eb:Messaging>
    <wsse:Security>
      <ds:Signature>        ← XML-DSIG over #body (RSA-SHA256)
        <ds:KeyInfo>
          <ds:X509Data><ds:X509Certificate>...</ds:X509Certificate></ds:X509Data>
        </ds:KeyInfo>
      </ds:Signature>
      <xenc:EncryptedKey>    ← RSA-OAEP wrapped AES-256 key
        <xenc:EncryptionMethod Algorithm="rsa-oaep-mgf1p"/>
        <ds:KeyInfo>
          <ds:X509Data><ds:X509Certificate>...</ds:X509Certificate></ds:X509Data>
        </ds:KeyInfo>
        <xenc:CipherData><xenc:CipherValue>...</xenc:CipherValue></xenc:CipherData>
      </xenc:EncryptedKey>
    </wsse:Security>
  </S12:Header>
  <S12:Body wsu:Id="body">
    <xenc:EncryptedData>   ← AES-256 encrypted UBL Invoice
      <xenc:CipherData><xenc:CipherValue>...</xenc:CipherValue></xenc:CipherData>
    </xenc:EncryptedData>
  </S12:Body>
</S12:Envelope>""")

doc.add_heading('4.10 PEPPOL Invitation Flow', level=2)
doc.add_paragraph(
    'For onboarding customers to PEPPOL, the platform supports a token-based invitation flow. '
    'An organisation admin sends an invitation to a customer, who completes registration '
    'via a secure link — no manual configuration of Access Points or participant links needed.'
)

add_api_endpoint(doc, 'POST', '/api/v1/my/invitations',
    'Org admin sends a PEPPOL invitation to a customer.')
add_api_endpoint(doc, 'GET', '/api/v1/invitations/{token}',
    'Validate invitation token (public).')
add_api_endpoint(doc, 'POST', '/api/v1/invitations/{token}/complete',
    'Buyer completes registration (public — token-authenticated).')

add_code_block(doc, """\
// Send invitation (org admin)
POST /api/v1/my/invitations
{
  "customerEmail": "buyer@acmecorp.co.zw",
  "participantId": "0190:ZW987654321",
  "message": "Please register for PEPPOL e-invoicing"
}

// Complete registration (buyer clicks link from email)
POST /api/v1/invitations/{token}/complete
{
  "endpointUrl": "https://erp.acmecorp.co.zw/peppol/receive",
  "simplifiedHttpDelivery": true,
  "deliveryAuthToken": "optional-bearer-token",
  "certificate": "-----BEGIN CERTIFICATE-----\\n..."
}""")

doc.add_paragraph(
    'When the buyer completes registration, the platform: '
    '(1) creates a RECEIVER AccessPoint for the buyer, '
    '(2) creates a PEPPOL participant link from the org to the buyer, '
    '(3) sends a confirmation email. '
    'Invitations expire after 72 hours.'
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 5 — ERP Adapters
# ═══════════════════════════════════════════════════════════════════════════════
doc.add_heading('5. ERP Adapters', level=1)

doc.add_heading('5.1 Architecture (Hexagonal Ports & Adapters)', level=2)
doc.add_paragraph(
    'ERP integration follows the hexagonal architecture pattern. An ErpInvoicePort interface '
    'defines the contract for fetching invoices and PDFs. Each ERP system has a dedicated adapter '
    'implementation. The ErpAdapterRegistry auto-discovers all adapters and indexes them by ErpSource.'
)

add_code_block(doc, """\
┌─────────────────────┐
│   ErpInvoicePort    │  ← Port (interface)
│   ──────────────    │
│  + supports()       │
│  + fetchInvoices()  │
│  + fetchInvoicePdf()│
│  + healthCheck()    │
└─────────────────────┘
         ▲
         │ implements
         │
┌────────┴────────┬──────────┬───────────┬──────────┐
│                 │          │           │          │
▼                 ▼          ▼           ▼          ▼
Sage         QuickBooks  Dynamics   Odoo     Generic
Intacct      Online      365        Adapter   API
Adapter      Adapter     Adapter              Adapter
(cond:       (cond:      (cond:     (cond:    (always)
SAGE_BASE    QB_CLIENT   D365_      ODOO_
_URL)        _ID)        BASE_URL)  BASE_URL)""")

doc.add_paragraph(
    'Each adapter is conditionally activated — set the base URL or client ID in the environment '
    'to enable it. The dispatch endpoint /api/v1/erp/dispatch accepts an erpSource parameter '
    'that selects which adapter to use.'
)

doc.add_paragraph('Available environment variables for ERP configuration:')
add_table(doc,
    ['ERP', 'Activation Env Var', 'Config Env Vars'],
    [
        ['Sage Intacct', 'SAGE_BASE_URL', 'SAGE_SENDER_ID, SAGE_SENDER_PASSWORD, SAGE_COMPANY_ID, SAGE_USER_ID, SAGE_USER_PASSWORD'],
        ['Sage Network', 'SAGE_NETWORK_BASE_URL', 'SAGE_NETWORK_API_KEY'],
        ['QuickBooks Online', 'QB_CLIENT_ID', 'QB_CLIENT_SECRET, QB_REALM_ID, QB_REFRESH_TOKEN'],
        ['Dynamics 365', 'D365_BASE_URL', 'D365_TENANT_ID, D365_CLIENT_ID, D365_CLIENT_SECRET'],
        ['Odoo', 'ODOO_BASE_URL', 'ODOO_DATABASE, ODOO_USERNAME, ODOO_API_KEY'],
    ])

doc.add_page_break()
doc.add_heading('5.2 Sage Intacct', level=2)
doc.add_paragraph(
    'The Sage Intacct adapter connects to Sage\'s XML API to fetch invoice data and PDFs.'
)
add_bullet(doc, 'Protocol: Sage XML API (readByQuery on ARINVOICE object)')
add_bullet(doc, 'Activation: Set SAGE_BASE_URL and related credentials')
add_bullet(doc, 'Invoice fields: RECORDNO → erpInvoiceId, DOCNUMBER → invoiceNumber, TOTALDUE → totalAmount, etc.')
add_bullet(doc, 'PDF source: Configured export directory (invoicePdfDir) or Sage report output')
add_bullet(doc, 'Auth: Sender ID + password + company + user credentials')

doc.add_page_break()
doc.add_heading('5.3 Sage Network', level=2)
doc.add_paragraph(
    'The Sage Network adapter synchronises PEPPOL delivery status with Sage\'s e-invoice '
    'lifecycle. It is NOT an ErpInvoicePort — it is a status sync and webhook handler.'
)
add_bullet(doc, 'Status lifecycle: New → Sending → Sent → Acknowledged / Rejected / Failed')
add_bullet(doc, 'Outbound: REST GET to Sage Network API to poll e-invoice status')
add_bullet(doc, 'Inbound: Webhook endpoint POST /webhooks/sage/einvoice-status/{orgId} receives Sage push notifications')
add_bullet(doc, 'Auth: API key in header')

doc.add_page_break()
doc.add_heading('5.4 QuickBooks Online', level=2)
doc.add_paragraph(
    'The QuickBooks Online adapter uses Intuit\'s REST API to fetch invoices and download PDFs.'
)
add_bullet(doc, 'Protocol: QBO REST API (GET /v3/company/{realmId}/invoice/{id})')
add_bullet(doc, 'Activation: Set QB_CLIENT_ID and related OAuth credentials')
add_bullet(doc, 'Invoice fields: Id → erpInvoiceId, DocNumber → invoiceNumber, TotalAmt → totalAmount, etc.')
add_bullet(doc, 'PDF source: QBO dedicated endpoint GET /invoice/{id}/pdf')
add_bullet(doc, 'Auth: OAuth 2.0 with refresh token rotation')

doc.add_page_break()
doc.add_heading('5.5 Microsoft Dynamics 365', level=2)
doc.add_paragraph(
    'The Dynamics 365 adapter connects to D365 Finance & Operations via OData API.'
)
add_bullet(doc, 'Protocol: D365 OData REST API (GET /data/SalesInvoiceHeadersV2)')
add_bullet(doc, 'Activation: Set D365_BASE_URL and related credentials')
add_bullet(doc, 'Invoice fields: InvoiceId → erpInvoiceId, InvoiceNumber → invoiceNumber, etc.')
add_bullet(doc, 'PDF source: Option A — export directory; Option B — Document Management / SharePoint attachments')
add_bullet(doc, 'Auth: OAuth 2.0 client credentials (tenant + client ID + client secret)')

doc.add_page_break()
doc.add_heading('5.6 Odoo', level=2)
doc.add_paragraph(
    'The Odoo adapter uses the JSON-RPC API to fetch invoices and generate PDF reports.'
)
add_bullet(doc, 'Protocol: Odoo JSON-RPC via /web/dataset/call_kw with search_read on account.move')
add_bullet(doc, 'Activation: Set ODOO_BASE_URL and related credentials')
add_bullet(doc, 'Invoice fields: name → invoiceNumber, invoice_date, amount_total, etc.')
add_bullet(doc, 'PDF source: Odoo report endpoint GET /report/pdf/account.report_invoice/{id}')
add_bullet(doc, 'Auth: API key (Odoo 14+)')

doc.add_page_break()
doc.add_heading('5.7 Generic API', level=2)
doc.add_paragraph(
    'The Generic API adapter is always active and serves as a fallback. It provides no '
    'ERP connectivity — invoice data must be provided directly in the API payload. '
    'This is used for manual dispatch from the frontend or external systems that send '
    'invoice data inline.'
)

add_api_endpoint(doc, 'GET', '/api/v1/erp/adapters',
    'List all registered and active ERP adapters.')
add_api_endpoint(doc, 'GET', '/api/v1/erp/health/{erpSource}',
    'Check connectivity to a specific ERP system.')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 6 — Billing & Metering
# ═══════════════════════════════════════════════════════════════════════════════
doc.add_heading('6. Billing & Metering', level=1)
doc.add_paragraph(
    'The platform uses a tiered rate-card billing model. Each organisation is assigned a '
    'RateProfile that defines a base monthly fee and per-invoice rates at different volume tiers.'
)

add_bullet(doc, 'RateProfile: Base monthly fee + multiple RateTier entries (volume range + rate per invoice)')
add_bullet(doc, 'Metering: Every delivery attempt (success or failure) is recorded as a UsageRecord')
add_bullet(doc, 'Billing Cycle: Monthly — period lifecycle OPEN → CLOSED → INVOICED → PAID')
add_bullet(doc, 'Auto-close: BillingScheduler runs at 00:05 on the 1st of each month')
add_bullet(doc, 'Platform Invoice: Generated via Thymeleaf template and emailed to org\'s accountsEmail')

doc.add_heading('Admin Billing Endpoints', level=2)
add_api_endpoint(doc, 'POST', '/api/v1/billing/rate-profiles', 'Create rate profile')
add_api_endpoint(doc, 'GET', '/api/v1/billing/rate-profiles', 'List rate profiles')
add_api_endpoint(doc, 'POST', '/api/v1/billing/estimate', 'Get cost estimate')
add_api_endpoint(doc, 'GET', '/api/v1/billing/summary/{orgId}/{period}', 'Get billing summary')
add_api_endpoint(doc, 'POST', '/api/v1/billing/close/{orgId}/{period}', 'Close billing period')
add_api_endpoint(doc, 'POST', '/api/v1/billing/invoice/{orgId}/{period}', 'Generate & send platform invoice')
add_api_endpoint(doc, 'POST', '/api/v1/billing/payment/{orgId}/{period}', 'Record payment')
add_api_endpoint(doc, 'GET', '/api/v1/billing/report/{period}', 'Revenue report')
add_api_endpoint(doc, 'GET', '/api/v1/billing/export/{orgId}/{period}', 'CSV export')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 7 — Webhooks
# ═══════════════════════════════════════════════════════════════════════════════
doc.add_heading('7. Webhooks & Callbacks', level=1)

doc.add_paragraph('The platform supports several webhook patterns:')

add_table(doc,
    ['Webhook', 'Type', 'Trigger', 'Signed'],
    [
        ['Campaign Completion', 'Outbound (POST to customer ERP)', 'Campaign finishes dispatching', 'HMAC-SHA256'],
        ['C4 PEPPOL Forwarding', 'Outbound (POST to org webhook)', 'Inbound PEPPOL document received', 'Bearer token'],
        ['Sage Network e-Invoice Status', 'Inbound (POST from Sage)', 'Sage e-invoice status change', 'TODO'],
    ])

doc.add_heading('Outbound Campaign Completion Webhook', level=2)
doc.add_paragraph(
    'Configured per-campaign via the callbackUrl field. Fires when all recipients '
    'have been processed (COMPLETED or PARTIALLY_FAILED). '
    'Signature: HMAC-SHA256(raw JSON body) using WEBHOOK_SECRET.'
)

doc.add_heading('C4 PEPPOL Forwarding', level=2)
doc.add_paragraph(
    'Configured per-organisation (c4WebhookUrl and c4WebhookAuthToken fields in the '
    'Organisation entity). Inbound PEPPOL documents are forwarded to this URL as raw UBL XML. '
    'Retries up to 3 times with 30-second intervals.'
)

doc.add_heading('Sage Network Inbound Webhook', level=2)
doc.add_paragraph(
    'Sage Network pushes e-invoice status changes to POST /webhooks/sage/einvoice-status/{orgId}. '
    'The platform updates the corresponding PeppolDeliveryRecord status based on the Sage lifecycle.'
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 8 — Error Handling & Retry
# ═══════════════════════════════════════════════════════════════════════════════
doc.add_heading('8. Error Handling & Retry', level=1)

doc.add_heading('Email Delivery', level=2)
add_bullet(doc, 'Retry: Up to 3 retries with exponential backoff (2s, 4s, 8s) via @Retryable')
add_bullet(doc, 'Fallback chain: Brevo API → Spring JavaMailSender SMTP')
add_bullet(doc, 'Rate limiting: Semaphore-based, configured via RATE_LIMIT (default 100 concurrent)')
add_bullet(doc, 'Batch recovery: Failed recipients can be retried via POST /api/v1/campaigns/{id}/retry')

doc.add_heading('PEPPOL Delivery', level=2)
add_bullet(doc, 'Schematron validation: Fatal violations prevent transmission, document marked FAILED')
add_bullet(doc, 'Network failures: Delivery record updated with error message, retryable flag set')
add_bullet(doc, 'Manual retry: POST /api/v1/dashboard/{orgId}/retry/{deliveryRecordId} resets FAILED to PENDING')

doc.add_heading('C4 Routing', level=2)
add_bullet(doc, 'Retry: Up to 3 attempts per document, 30 seconds between retries')
add_bullet(doc, 'No webhook URL: Document skipped without retry increment')
add_bullet(doc, '2xx required: Any non-2xx response increments retry count')
add_bullet(doc, 'Final action: After 3 failures, document marked ROUTING_FAILED')

doc.add_heading('API Error Response Codes', level=2)
add_table(doc,
    ['HTTP Status', 'Meaning', 'When'],
    [
        ['200', 'Success', 'Request completed normally'],
        ['201', 'Created', 'Resource created (campaign, AP, link)'],
        ['400', 'Bad Request', 'Validation error, invalid payload'],
        ['401', 'Unauthorized', 'Missing or invalid API key / session'],
        ['403', 'Forbidden', 'Not authorized for this resource'],
        ['404', 'Not Found', 'Resource not found'],
        ['409', 'Conflict', 'Duplicate participant ID or payload'],
        ['502', 'Bad Gateway', 'SMTP / ERP / AS4 upstream failure'],
    ])

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 9 — Security
# ═══════════════════════════════════════════════════════════════════════════════
doc.add_heading('9. Security Overview', level=1)

doc.add_heading('Authentication', level=2)
add_table(doc,
    ['Mechanism', 'Scope', 'Header'],
    [
        ['API Key', 'Organisation-scoped', 'X-API-Key'],
        ['Admin Session', 'Full admin access', 'Authorization: Bearer {token}'],
        ['Org Session', 'Org-scoped', 'Authorization: Bearer {token}'],
        ['Invitation Token', 'Single-use registration', 'URL path parameter'],
    ])

doc.add_heading('PEPPOL Security', level=2)
add_bullet(doc, 'AS4 Inbound: HMAC-SHA256 with pre-shared secret (constant-time comparison)')
add_bullet(doc, 'AS4 Outbound: XML-DSIG (RSA-SHA256) + XML-Enc (AES-256 / RSA-OAEP)')
add_bullet(doc, 'Deduplication: SHA-256 payload hash prevents replay attacks')
add_bullet(doc, 'Fail-closed: No shared secret = all inbound traffic rejected')
add_bullet(doc, 'Certificate rotation: Current certificate can be rotated without downtime')

doc.add_heading('Webhook Security', level=2)
add_bullet(doc, 'Outbound campaign webhooks: HMAC-SHA256 signed with configurable secret (>=32 chars)')
add_bullet(doc, 'C4 forwarding: Optional Bearer token authentication per organisation')

doc.add_heading('Data Protection', level=2)
add_bullet(doc, 'PDF size cap: 10 MB default, configurable via MAX_PDF_BYTES (OOM protection)')
add_bullet(doc, 'Path traversal prevention: File-path PDFs must be under configured base path')
add_bullet(doc, 'AS4 messages: End-to-end encrypted (XML-Enc AES-256)')
add_bullet(doc, 'Database: Connection-level TLS recommended')

doc.add_heading('Public Endpoints', level=2)
doc.add_paragraph(
    'The following endpoints are publicly accessible (no authentication required):'
)
add_table(doc,
    ['Endpoint', 'Purpose', 'Auth Method'],
    [
        ['POST /peppol/as4/receive', 'PEPPOL C3 inbound', 'HMAC-SHA256 (in-app)'],
        ['GET /peppol/as4/health', 'PEPPOL AS4 health check', 'None'],
        ['GET /bdxr/smp/**', 'SMP service metadata', 'None'],
        ['POST /webhooks/sage/**', 'Sage Network webhooks', 'HMAC (planned)'],
        ['POST /api/v1/admin/login', 'Admin login', 'Password'],
        ['POST /api/v1/org/login', 'Org login', 'Password'],
        ['GET /api/v1/invitations/{token}', 'Validate invitation', 'Token'],
        ['POST /api/v1/invitations/{token}/complete', 'Complete invitation', 'Token'],
    ])

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 10 — Deployment
# ═══════════════════════════════════════════════════════════════════════════════
doc.add_heading('10. Deployment', level=1)

doc.add_heading('Prerequisites', level=2)
add_bullet(doc, 'Java 25 or later')
add_bullet(doc, 'PostgreSQL 16+')
add_bullet(doc, 'Maven 3.9+')
add_bullet(doc, 'Node.js 20+ / npm (for frontend build)')
add_bullet(doc, 'nginx (for production reverse proxy)')
add_bullet(doc, 'Let\'s Encrypt or equivalent TLS certificates')

doc.add_heading('Required Environment Variables', level=2)
doc.add_paragraph('These must be set for the application to start:')
add_table(doc,
    ['Variable', 'Description', 'Example'],
    [
        ['DB_PASS', 'PostgreSQL password', 'strong-password'],
        ['ADMIN_PASSWORD', 'Admin login password', 'strong-password'],
        ['WEBHOOK_SECRET', 'HMAC signing key (>=32 chars)', 'a-string-of-at-least-32-characters'],
        ['SPRING_MAIL_PASSWORD', 'SMTP/Brevo password', 'smtp-password'],
    ])

doc.add_paragraph('Deployment scripts are available in the repository root:')
add_bullet(doc, 'deploy-native.sh — bare-metal deployment (systemd + nginx)')
add_bullet(doc, 'deploy.sh — Docker Compose deployment')

doc.add_heading('Health Checks', level=2)
add_api_endpoint(doc, 'GET', '/actuator/health', 'Overall health — includes database, mail, PEPPOL.')
add_api_endpoint(doc, 'GET', '/actuator/info', 'Build and version metadata.')
add_api_endpoint(doc, 'GET', '/peppol/as4/health', 'PEPPOL AS4 endpoint health.')

doc.add_paragraph()
doc.add_paragraph()

# ── Footer ───────────────────────────────────────────────────────────────────
footer = doc.add_paragraph()
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = footer.add_run('─' * 60)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('InvoiceDirect Platform — Developer Integration Guide v1.0')
run.font.size = Pt(8)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(f'Generated {datetime.date.today().isoformat()} — eSolutions (Pvt) Ltd')
run.font.size = Pt(8)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

# ── Save ─────────────────────────────────────────────────────────────────────
output_path = '/Users/luckyncube/Library/CloudStorage/OneDrive-Personal/dev/mass-mailer/InvoiceDirect-Developer-Guide.docx'
doc.save(output_path)
print(f'Document saved to: {output_path}')
print(f'Total sections: 10')
