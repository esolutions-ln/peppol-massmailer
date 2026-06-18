#!/usr/bin/env python3
"""Generate UAT documents for the InvoiceDirect Mass Mailer system."""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
import datetime

OUTPUT_DIR = "/Users/luckyncube/Library/CloudStorage/OneDrive-Personal/dev/mass-mailer/uat-documents"

def set_cell_shading(cell, color):
    """Set background shading on a table cell."""
    shading = cell._element.get_or_add_tcPr()
    shading_elem = shading.makeelement(qn('w:shd'), {
        qn('w:fill'): color,
        qn('w:val'): 'clear'
    })
    shading.append(shading_elem)

def add_header_row(table, headers, color="1F4E79"):
    """Style the first row of a table as a header."""
    for i, text in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = text
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.size = Pt(9)
        set_cell_shading(cell, color)

def add_data_row(table, values, bold_first=False):
    """Add a data row to a table."""
    row = table.add_row()
    for i, text in enumerate(values):
        cell = row.cells[i]
        cell.text = str(text)
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(9)
                if bold_first and i == 0:
                    run.bold = True
    return row

def style_document(doc):
    """Apply base styling to the document."""
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(10)

    # Set narrow margins
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

def add_heading(doc, text, level=1):
    """Add a styled heading."""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        if level == 1:
            run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
        elif level == 2:
            run.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)
        elif level == 3:
            run.font.color.rgb = RGBColor(0x44, 0x72, 0xC4)
    return heading

# ============================================================
# DOCUMENT 1: UAT TEST PLAN
# ============================================================

def generate_test_plan():
    doc = Document()
    style_document(doc)

    # Title page
    title = doc.add_heading('User Acceptance Test Plan', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('InvoiceDirect Mass Mailer System')
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(f'\nVersion: 2.0\nDate: {datetime.date.today().strftime("%d %B %Y")}\nStatus: Draft').font.size = Pt(11)

    doc.add_page_break()

    # Table of Contents placeholder
    add_heading(doc, 'Table of Contents', level=1)
    doc.add_paragraph('1. Introduction')
    doc.add_paragraph('2. Scope')
    doc.add_paragraph('3. Test Objectives')
    doc.add_paragraph('4. Features to Be Tested')
    doc.add_paragraph('5. Features Not to Be Tested')
    doc.add_paragraph('6. Test Approach')
    doc.add_paragraph('7. Entry and Exit Criteria')
    doc.add_paragraph('8. Test Schedule')
    doc.add_paragraph('9. Roles and Responsibilities')
    doc.add_paragraph('10. Test Deliverables')
    doc.add_paragraph('11. Risk and Mitigation')
    doc.add_paragraph('12. Approval')

    doc.add_page_break()

    # 1. Introduction
    add_heading(doc, '1. Introduction', level=1)
    doc.add_paragraph(
        'This plan covers User Acceptance Testing (UAT) of the InvoiceDirect Mass Mailer. '
        'The release under test focuses on the four core business scenarios:'
    )
    intro_items = [
        'Capture or import customers into the platform',
        'Receive fiscalised invoice PDFs into a watched folder',
        'Pick a PDF invoice and email it to the recipient',
        'Enforce that the recipient is an existing customer on the platform'
    ]
    for item in intro_items:
        doc.add_paragraph(item, style='List Bullet')
    doc.add_paragraph(
        'A small set of optimisations that the code already supports — auto-CC of sibling contacts, '
        'customisable email body templates, and an in-app PDF viewer for sent invoices — are included '
        'as confidence-builders rather than expansion of scope.'
    )

    # 2. Scope
    add_heading(doc, '2. Scope', level=1)
    doc.add_paragraph('In-scope functional areas:')
    scope_items = [
        'F-01 — Customer Capture (single): Create a single customer with one or more contacts via the UI or API.',
        'F-02 — Customer Bulk Import (CSV): Upload a CSV; importer upserts customers by erpCustomerId and links contacts.',
        'F-03 — Fiscalised Invoices in Folder: Drop ZIMRA-signed PDFs (and optional sidecar JSON) into the watched inbox; the watcher ingests them automatically.',
        'F-04 — Pick PDF & Send Email: From the invoices list (or via API) select a PDF and email it to the buyer.',
        'F-05 — Customer Must Be On Platform: Recipient must resolve to an existing customer (auto-upserted from the PDF Buyer block or pre-loaded).',
        'F-06 — Custom Email Body Template: Per-organisation editable template with {{placeholders}}; default applies automatically, optional per-send override.',
        'F-07 — Auto-CC Sibling Contacts: When sending to one contact of a customer, all other contacts of the same customer are added as Cc.',
        'F-08 — View Sent Invoice PDF: From the Invoices page, open and download the PDF that was actually delivered.',
        'F-09 — Email Delivery Monitoring: Campaign status tracking, MailRecipient status transitions, and retry of failed deliveries.',
    ]
    for item in scope_items:
        doc.add_paragraph(item, style='List Bullet')

    # 3. Test Objectives
    add_heading(doc, '3. Test Objectives', level=1)
    objectives = [
        'Confirm an admin can capture a customer manually and via CSV import.',
        'Confirm the PDF watcher ingests fiscalised invoices dropped into the inbox folder.',
        'Confirm a user can select a stored PDF invoice and email it to the recipient.',
        'Confirm dispatch is rejected (or auto-creates the customer from the Buyer block) when the recipient is not yet on the platform.',
        'Confirm the default custom email template is applied; an override at send time replaces it.',
        'Confirm all sibling contacts of the same customer are CC\'d automatically.',
        'Confirm the sent PDF can be viewed and downloaded from the Invoices page.',
        'Confirm MailRecipient status transitions correctly (PENDING → SENT/FAILED) and failed recipients can be retried.',
    ]
    for obj in objectives:
        doc.add_paragraph(obj, style='List Bullet')

    # 4. Features to Be Tested
    add_heading(doc, '4. Features to Be Tested', level=1)
    features = [
        ('F-01', 'Customer Capture (single)', 'Create a single customer with one or more contacts via the UI or API.'),
        ('F-02', 'Customer Bulk Import (CSV)', 'Upload a CSV; importer upserts customers by erpCustomerId and links contacts.'),
        ('F-03', 'Fiscalised Invoices in Folder', 'Drop ZIMRA-signed PDFs (and optional sidecar JSON) into the watched inbox; the watcher ingests them automatically. Includes ZIMRA FDMS fiscal marker validation.'),
        ('F-04', 'Pick PDF & Send Email', 'From the invoices list (or via API) select a PDF and email it to the buyer.'),
        ('F-05', 'Customer Must Be On Platform', 'Recipient must resolve to an existing customer (auto-upserted from the PDF Buyer block or pre-loaded).'),
        ('F-06', 'Custom Email Body Template', 'Per-organisation editable template with {{placeholders}}; default applies automatically, optional per-send override.'),
        ('F-07', 'Auto-CC Sibling Contacts', 'When sending to one contact of a customer, all other contacts of the same customer are added as Cc.'),
        ('F-08', 'View Sent Invoice PDF', 'From the Invoices page, open and download the PDF that was actually delivered.'),
        ('F-09', 'Email Delivery Monitoring', 'Campaign status lifecycle, MailRecipient status transitions, and retry of failed recipients.'),
    ]

    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_header_row(table, ['Ref', 'Feature', 'Description'])
    for ref, feat, desc in features:
        row = add_data_row(table, [ref, feat, desc])
        set_cell_shading(row.cells[0], "E8F0FE")
        row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # 5. Features Not to Be Tested
    add_heading(doc, '5. Features Not to Be Tested', level=1)
    not_tested = [
        'ERP adapter end-to-end integration (Sage, QuickBooks, Odoo, Dynamics 365) — covered by adapter-level tests.',
        'PEPPOL/AS4 outbound delivery — verified separately in the access-point test plan.',
        'Platform billing and invoicing of organisations — separate workstream.',
        'Non-functional concerns (scaling, DR, infra hardening).'
    ]
    for item in not_tested:
        doc.add_paragraph(item, style='List Bullet')

    # 6. Test Approach
    add_heading(doc, '6. Test Approach', level=1)
    doc.add_paragraph(
        'UAT runs against the staging environment using anonymised production-like data. '
        'Each test case is executed manually and the outcome (Pass / Fail / Blocked) recorded '
        'with actual results and evidence.'
    )
    doc.add_paragraph('Test execution follows these phases:')
    phases = [
        'Phase 1 — Environment readiness: backend port 9199 and frontend port 8199 healthy; database reachable.',
        'Phase 2 — Customer onboarding: F-01 and F-02 executed end-to-end.',
        'Phase 3 — Invoice ingestion and dispatch: F-03, F-04, F-05 executed end-to-end.',
        'Phase 4 — Optimisations: F-06, F-07, F-08.',
        'Phase 5 — Sign-off: collate results, raise defects, obtain stakeholder approval.',
    ]
    for phase in phases:
        doc.add_paragraph(phase)

    # 7. Entry and Exit Criteria
    add_heading(doc, '7. Entry and Exit Criteria', level=1)
    add_heading(doc, '7.1 Entry Criteria', level=2)
    entry = [
        'Staging environment deployed and healthy (/actuator/health returns UP).',
        'Inbox folder exists, is writable by the backend container, and the watcher is enabled (massmailer.pdf-watcher.enabled=true).',
        'Default organisation id is configured for the watcher OR a sidecar JSON is supplied per file.',
        'SMTP credentials are set; a test mailbox is available for delivery verification.',
        'All P0/P1 defects from prior test cycles are resolved.',
    ]
    for item in entry:
        doc.add_paragraph(item, style='List Bullet')

    add_heading(doc, '7.2 Exit Criteria', level=2)
    exit_criteria = [
        'Every test case in document 02 has been executed.',
        'Pass rate >= 95%.',
        'No open P0 or P1 defects.',
        'Sign-off document (03) is approved by all named stakeholders.',
    ]
    for item in exit_criteria:
        doc.add_paragraph(item, style='List Bullet')

    # 8. Test Schedule
    add_heading(doc, '8. Test Schedule', level=1)
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_header_row(table, ['Activity', 'Duration', 'Start Date', 'End Date'])
    schedule = [
        ['Environment Setup & Verification', '0.5 day', 'TBD', 'TBD'],
        ['Customer Onboarding (F-01, F-02)', '1 day', 'TBD', 'TBD'],
        ['Invoice Ingestion & Dispatch (F-03, F-04, F-05)', '1 day', 'TBD', 'TBD'],
        ['Optimisations & Monitoring (F-06, F-07, F-08, F-09)', '0.5 day', 'TBD', 'TBD'],
        ['Defect Retesting & Sign-off', '0.5 day', 'TBD', 'TBD'],
    ]
    for row_data in schedule:
        add_data_row(table, row_data)

    doc.add_paragraph()

    # 9. Roles and Responsibilities
    add_heading(doc, '9. Roles and Responsibilities', level=1)
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_header_row(table, ['Role', 'Name', 'Responsibility'])
    roles = [
        ['Test Coordinator', 'TBD', 'Plan, coordinate, and report UAT activities'],
        ['Business UAT Tester', 'TBD', 'Execute test cases, document results, raise defects'],
        ['Technical Lead', 'TBD', 'Resolve environment issues, review defects'],
        ['Project Stakeholder', 'TBD', 'Review results, approve sign-off'],
    ]
    for row_data in roles:
        add_data_row(table, row_data)

    doc.add_paragraph()

    # 10. Test Deliverables
    add_heading(doc, '10. Test Deliverables', level=1)
    deliverables = [
        'UAT Test Plan (this document)',
        'UAT Test Cases document with execution results',
        'Defect log (raised in issue tracker)',
        'UAT Sign-off document'
    ]
    for item in deliverables:
        doc.add_paragraph(item, style='List Bullet')

    # 11. Risk and Mitigation
    add_heading(doc, '11. Risk and Mitigation', level=1)
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_header_row(table, ['Risk', 'Impact', 'Mitigation'])
    risks = [
        ['Environment instability or configuration drift', 'Delays in testing', 'Use infrastructure-as-code; verify environment readiness before testing'],
        ['Data migration issues discovered during testing', 'Data integrity concerns', 'Run migration verification scripts before functional testing'],
        ['Limited availability of business testers', 'Incomplete test coverage', 'Schedule testing sessions in advance; document hand-off notes'],
        ['Integration dependencies unavailable (ERP, PEPPOL)', 'Cannot test end-to-end flow', 'Use mock/stub adapters for isolated testing'],
    ]
    for row_data in risks:
        add_data_row(table, row_data)

    doc.add_paragraph()

    # 12. Approval
    add_heading(doc, '12. Approval', level=1)
    doc.add_paragraph('The following stakeholders have reviewed and approved this UAT Test Plan:')
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_header_row(table, ['Role', 'Name', 'Signature', 'Date'])
    for _ in range(3):
        add_data_row(table, ['', '', '', ''])

    doc.save(f'{OUTPUT_DIR}/01-UAT-Test-Plan.docx')
    print(f"✓ Generated: 01-UAT-Test-Plan.docx")


# ============================================================
# DOCUMENT 2: UAT TEST CASES
# ============================================================

def generate_test_cases():
    doc = Document()
    style_document(doc)

    title = doc.add_heading('User Acceptance Test Cases', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('InvoiceDirect Mass Mailer System')
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(f'\nVersion: 2.0\nDate: {datetime.date.today().strftime("%d %B %Y")}\nTotal Test Cases: 22').font.size = Pt(11)

    doc.add_page_break()

    # Legend / Status Key
    add_heading(doc, 'Status Key', level=1)
    key_table = doc.add_table(rows=1, cols=2)
    key_table.style = 'Table Grid'
    add_header_row(key_table, ['Status', 'Meaning'])
    add_data_row(key_table, ['Pass', 'Actual result matches expected result'])
    add_data_row(key_table, ['Fail', 'Actual result does not match expected result'])
    add_data_row(key_table, ['Blocked', 'Could not execute due to dependency or environment issue'])
    add_data_row(key_table, ['N/A', 'Not applicable for this test cycle'])
    doc.add_paragraph()

    # All test cases organized by feature
    test_cases = [
        # ============================================================
        # F-01: Customer Capture (single) — via UI or API
        # ============================================================
        ('F-01', 'TC-01', 'Create single customer with one contact via UI', 'Critical',
         'As an admin user, I can create a single customer with one contact via the Customers page.',
         '1. Open the Customers page (admin view)\n2. Click "New Customer"\n3. Fill companyName, erpCustomerId, primary contact email and name\n4. Save',
         'Customer appears at the top of the list.\nDetail panel shows the contact and an erpSource value.\nPOST /api/v1/organizations/{orgId}/customers returned 200.',
         'Pass / Fail / Blocked', ''),

        ('F-01', 'TC-02', 'Create customer with multiple contacts via API', 'High',
         'As an API consumer, I can create a customer with multiple contacts in one call.',
         '1. POST /api/v1/organizations/{orgId}/customers with a contacts array of two entries\n2. Verify HTTP 200',
         'One customer row and two contact rows are created.\nBoth contacts are linked to the same customer.\nDuplicate erpCustomerId is rejected with HTTP 409 on a second call.',
         'Pass / Fail / Blocked', ''),

        # ============================================================
        # F-02: Customer Bulk Import (CSV)
        # ============================================================
        ('F-02', 'TC-03', 'CSV import with auto-suggested column mapping', 'Critical',
         'As an admin user, I can import customers from a CSV with custom column mapping.',
         '1. From the Customer Import dialog, upload a CSV\n2. Confirm the auto-suggested mapping and adjust if necessary (e.g. FileReference → tinNumber)\n3. Submit',
         'Response includes totalRows, created, updated, skipped, errors=0.\nCustomers tab now shows the imported rows.\nRe-running the same CSV results in "updated" rather than "created".',
         'Pass / Fail / Blocked', ''),

        ('F-02', 'TC-04', 'CSV import upserts by erpCustomerId', 'High',
         'As a business user, the importer treats erpCustomerId as the unique key (upsert).',
         '1. Import a CSV with erpCustomerId=TEST-01\n2. Edit the company name in the CSV\n3. Re-import',
         'Only one customer row for TEST-01 exists.\nCompany name reflects the latest import.\nContacts on the row are upserted by email.',
         'Pass / Fail / Blocked', ''),

        ('F-02', 'TC-05', 'CSV import with no email column', 'Medium',
         'As an admin user, I can import a CSV with no email column without error.',
         '1. Import a CSV that has erpCustomerId and companyName but no email column',
         'Customer rows are created.\nNo contact rows are created for rows without an email.\nerrors=0 in the response.',
         'Pass / Fail / Blocked', ''),

        # ============================================================
        # F-03: Fiscalised Invoices in Folder (PDF Watcher)
        # ============================================================
        ('F-03', 'TC-06', 'PDF watcher ingests fiscalised PDF from inbox', 'Critical',
         'As a system, the PDF watcher ingests a fiscalised PDF dropped into the inbox folder.',
         '1. Drop a valid ZIMRA-signed PDF into the configured inbox directory\n2. Optionally drop a matching sidecar JSON with metadata\n3. Wait one poll interval',
         'Watcher log shows the file was picked up.\nAn invoice record is created for the configured default organisation.\nThe PDF is moved to the "emailed" directory once dispatch succeeds.',
         'Pass / Fail / Blocked', ''),

        ('F-03', 'TC-07', 'PDF watcher auto-upserts customer from Buyer block', 'High',
         'As a system, the watcher auto-upserts the customer from the PDF Buyer block.',
         '1. Drop a PDF whose buyer is not yet on the platform\n2. Allow the watcher to process it',
         'A new customer row is created (or an existing one is updated) for that buyer.\nerpSource is set on the customer.\nThe customer appears immediately in the Customers page.',
         'Pass / Fail / Blocked', ''),

        # ============================================================
        # F-04: Pick PDF & Send Email
        # ============================================================
        ('F-04', 'TC-08', 'Resend a stored PDF invoice to recipient', 'High',
         'As a user, I can pick a stored PDF invoice and re-send it to the recipient.',
         '1. Open the Invoices page\n2. Find the desired invoice and click "View"\n3. Click "Resend" (or invoke POST /api/v1/mail/invoice with the same payload)',
         'Recipient receives the email with the PDF attached.\nResponse is delivered=true with a messageId.\nThe new send is logged against the same invoice number.',
         'Pass / Fail / Blocked', ''),

        ('F-04', 'TC-09', 'One-off invoice send via PDF upload API', 'Critical',
         'As a user, I can send a one-off invoice by uploading a PDF via the API.',
         '1. POST /api/v1/mail/invoice/upload with the PDF as multipart and JSON metadata\n2. Verify HTTP 200',
         'Recipient receives the email.\nInvoice number, total, and currency from metadata are rendered in the body.\nThe attachment opens correctly.',
         'Pass / Fail / Blocked', ''),

        ('F-04', 'TC-10', 'Invalid email address rejected at validation', 'Medium',
         'As a user, sending with an invalid email address is rejected at validation.',
         '1. POST /api/v1/mail/invoice with to="not-an-email"',
         'HTTP 400 with a "well-formed email address" validation error.\nNo outbound SMTP attempt is made.',
         'Pass / Fail / Blocked', ''),

        # ============================================================
        # F-05: Customer Must Be On Platform (auto-upsert from Buyer block)
        # ============================================================
        ('F-05', 'TC-11', 'Watcher auto-registers unknown buyer before dispatch', 'Critical',
         'As a system, dispatch via the watcher auto-registers an unknown buyer before sending.',
         '1. Confirm the buyer email is NOT in the customers table\n2. Drop a PDF for that buyer into the inbox\n3. After processing, query the customer record',
         'A customer is now present for the buyer (auto-upserted from the Buyer block).\nThe invoice is dispatched successfully.\nPer memory rule: customers are never pre-seeded manually for this path.',
         'Pass / Fail / Blocked', ''),

        ('F-05', 'TC-12', 'Manual send to existing customer resolves for CC', 'High',
         'As a system, a manual send to a recipient already on the platform resolves the customer for CC purposes.',
         '1. Ensure the recipient is a contact under a known customer\n2. Send a single invoice to that recipient via the API',
         'Send succeeds.\nLog shows the customer was looked up and any sibling contacts attached as Cc.\nThe recipient\'s customer row\'s lastInvoiceSent timestamp is updated.',
         'Pass / Fail / Blocked', ''),

        # ============================================================
        # F-06: Custom Email Body Template
        # ============================================================
        ('F-06', 'TC-13', 'Create custom email template with placeholders', 'Medium',
         'As an organisation admin, I can edit a custom email body template with {{placeholders}}.',
         '1. Open Email Templates\n2. Create a template with subject "Invoice {{invoiceNumber}}" and body containing {{recipientName}}, {{totalAmount}}, {{currency}}\n3. Mark it as default and save',
         'Template is saved and marked default in the list.\nRe-opening the editor preserves the body verbatim.\nOnly one default template exists per organisation.',
         'Pass / Fail / Blocked', ''),

        ('F-06', 'TC-14', 'Default template auto-applied and overridable per send', 'Medium',
         'As a sender, the org default template is applied automatically and overridable per send.',
         '1. Send an invoice without specifying emailTemplateId\n2. Inspect the rendered email body\n3. Send again with emailTemplateId set to a different template',
         'First send uses the default template (placeholders substituted).\nSecond send uses the overriding template.\nIf no DB template exists the legacy Thymeleaf "invoice" template is used.',
         'Pass / Fail / Blocked', ''),

        # ============================================================
        # F-07: Auto-CC Sibling Contacts
        # ============================================================
        ('F-07', 'TC-15', 'Sibling contacts CC\'d on invoice send', 'High',
         'As a business user, sibling contacts of the same customer are CC\'d on invoice sends.',
         '1. Identify a customer with two contacts (A and B)\n2. Send an invoice TO contact A\'s email',
         'Contact A receives the email as "To".\nContact B receives the email as "Cc" with the same PDF.\nA customer with a single contact produces no Cc header.',
         'Pass / Fail / Blocked', ''),

        # ============================================================
        # F-08: View Sent Invoice PDF
        # ============================================================
        ('F-08', 'TC-16', 'View and download sent invoice PDF', 'Medium',
         'As a user, I can view and download the PDF of a previously sent invoice.',
         '1. Open the Invoices page\n2. Click "View" on a row marked SENT\n3. Inside the modal, click "Download PDF"',
         'PDF renders in the in-app viewer.\nDownloaded file matches the attachment that was emailed.\nPagination and search across invoices both work alongside the viewer.',
         'Pass / Fail / Blocked', ''),

        # ============================================================
        # F-03 (extended): ZIMRA FDMS Validation — NEW
        # ============================================================
        ('F-03', 'TC-17', 'ZIMRA FDMS fiscal markers validated on PDF watcher pickup', 'Critical',
         'As a system, the PDF watcher must validate that every processed PDF contains ZIMRA FDMS fiscal markers before dispatch.',
         '1. Drop a PDF containing valid ZIMRA FDMS markers (fdms.zimra.co.zw domain string, 16-char hex verification code, and fiscal device field) + sidecar JSON\n2. Wait for watcher to process\n3. Check application logs',
         'Watcher log confirms fiscal validation passed: "All ZIMRA fiscal markers present".\nPDF contains: (a) "fdms.zimra.co.zw" domain, (b) "Verification Code" label or 16-char uppercase hex pattern [0-9A-F]{16}, (c) at least one of "Device ID" / "Fiscal Day" / "Fiscal Invoice" / "Global Receipt".\nInvoice is dispatched successfully.',
         'Pass / Fail / Blocked',
         'ZimraFiscalValidator checks three layers: FDMS domain, verification code, and fiscal device field. All three must pass.'),

        ('F-03', 'TC-18', 'PDF missing fiscal markers rejected by watcher', 'Critical',
         'As a system, the watcher must reject PDFs that are missing ZIMRA FDMS fiscal markers.',
         '1. Generate a PDF that is a valid PDF (%PDF- magic bytes) but contains NONE of the ZIMRA FDMS markers (no fdms.zimra.co.zw, no verification code, no device ID)\n2. Drop it into the inbox with a sidecar JSON\n3. Wait for processing',
         'Watcher log shows fiscal validation FAILED with specific error messages.\nPDF is moved to inbox/failed/ directory (not emailed/).\nNo campaign or MailRecipient record is created for this invoice.\nError is logged with details of which markers are missing.',
         'Pass / Fail / Blocked',
         'Use pdfbox or similar to strip fiscal markers from a known-good PDF for testing. The three validation layers are: domain fdms.zimra.co.zw, verification code, and fiscal device field.'),

        # ============================================================
        # F-03 (extended): Status Tracking — NEW
        # ============================================================
        ('F-03', 'TC-19', 'MailRecipient status transitions to SENT after successful dispatch', 'Critical',
         'As a system, the MailRecipient status must reflect SENT/DELIVERED after email dispatch succeeds.',
         '1. Drop a valid fiscalised PDF + sidecar into the inbox\n2. Wait for watcher to process\n3. Query MailRecipient via GET /api/v1/campaigns/{campaignId}\n4. Query the database directly: SELECT status FROM mail_recipient WHERE invoice_number = ...',
         'MailRecipient.status = "SENT" or "DELIVERED".\nmessageId is populated (not null).\nsentAt timestamp is populated.\nattachmentSizeBytes matches the original PDF size.\nCampaign status = "COMPLETED".',
         'Pass / Fail / Blocked',
         'The statuses are: PENDING → SENT/DELIVERED / FAILED / SKIPPED. Check both the API response and the database directly.'),

        # ============================================================
        # F-09: Email Delivery Monitoring — NEW
        # ============================================================
        ('F-09', 'TC-20', 'Campaign status shows COMPLETED after batch dispatch', 'High',
         'As a user, the campaign status must reflect COMPLETED when all recipients have been processed.',
         '1. Drop 3 valid fiscalised PDFs into the inbox (same batch)\n2. Poll GET /api/v1/campaigns/{id} until processing completes\n3. Verify campaign status',
         'Campaign status transitions: CREATED → QUEUED → IN_PROGRESS → COMPLETED.\nsentCount = 3, failedCount = 0, skippedCount = 0.\ncompletedAt timestamp is populated.\nstartedAt is before completedAt.',
         'Pass / Fail / Blocked',
         'Campaign lifecycle: CREATED → QUEUED → IN_PROGRESS → COMPLETED / PARTIALLY_FAILED / FAILED.'),

        ('F-09', 'TC-21', 'Failed email delivery detected and status set to FAILED', 'High',
         'As a system, when email delivery fails (invalid recipient, SMTP error), the status must reflect FAILED.',
         '1. Drop a PDF with a deliberately invalid recipient email (e.g., nonexistent@domain.com) + sidecar\n2. Wait for processing\n3. Check campaign status\n4. Query MailRecipient',
         'MailRecipient.status = "FAILED".\nerrorMessage describes the failure reason (SMTP error / invalid address).\nretryCount is 0 (first attempt).\nCampaign status = "PARTIALLY_FAILED" (if some succeeded) or "FAILED" (if all failed).',
         'Pass / Fail / Blocked',
         'To test, configure the SMTP to a blackhole or use a known-bad recipient domain.'),

        ('F-09', 'TC-22', 'Retry failed recipients resets status to PENDING', 'Medium',
         'As a user, I can retry failed recipients and their status resets to PENDING for re-dispatch.',
         '1. Cause a delivery failure (see TC-21)\n2. POST /api/v1/campaigns/{id}/retry\n3. Poll campaign status\n4. Query MailRecipient for the retried record',
         'On retry, failed MailRecipient rows are reset to PENDING.\nThe dispatch pipeline re-runs for those recipients.\nIf the retry succeeds, status becomes SENT and retryCount is incremented.\nRetry respects maxRetries (default 3); after exceeding max, status stays FAILED.',
         'Pass / Fail / Blocked',
         'POST /api/v1/campaigns/{id}/retry triggers async retry via @Async("mailExecutor").'),
    ]

    # Group test cases by feature
    current_feature = None
    for ref, tc_id, title, priority, desc, steps, expected, result, notes in test_cases:
        if current_feature != ref:
            if current_feature is not None:
                doc.add_paragraph()  # spacing
            current_feature = ref
            feature_name = {
                'F-01': 'Customer Capture (single)',
                'F-02': 'Customer Bulk Import (CSV)',
                'F-03': 'Fiscalised Invoices in Folder',
                'F-04': 'Pick PDF & Send Email',
                'F-05': 'Customer Must Be On Platform',
                'F-06': 'Custom Email Body Template',
                'F-07': 'Auto-CC Sibling Contacts',
                'F-08': 'View Sent Invoice PDF',
                'F-09': 'Email Delivery Monitoring'
            }.get(ref, '')
            add_heading(doc, f'{ref}: {feature_name}', level=1)

        add_heading(doc, f'{tc_id}: {title}', level=2)

        # Meta info in a compact table
        meta_table = doc.add_table(rows=2, cols=4)
        meta_table.style = 'Table Grid'
        meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        # Row 1
        meta_table.rows[0].cells[0].text = 'Priority'
        meta_table.rows[0].cells[1].text = priority
        meta_table.rows[0].cells[2].text = 'Feature Ref'
        meta_table.rows[0].cells[3].text = ref
        for cell in meta_table.rows[0].cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
                    run.font.size = Pt(9)
        set_cell_shading(meta_table.rows[0].cells[0], "1F4E79")
        set_cell_shading(meta_table.rows[0].cells[2], "1F4E79")
        for paragraph in meta_table.rows[0].cells[0].paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        for paragraph in meta_table.rows[0].cells[2].paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        # Row 2
        meta_table.rows[1].cells[0].text = 'Description'
        meta_table.rows[1].cells[1].text = desc
        meta_table.rows[1].cells[2].text = 'Status'
        meta_table.rows[1].cells[3].text = result
        set_cell_shading(meta_table.rows[1].cells[0], "E8F0FE")
        set_cell_shading(meta_table.rows[1].cells[2], "E8F0FE")
        meta_table.rows[1].cells[1].paragraphs[0].runs[0].font.size = Pt(9)
        meta_table.rows[1].cells[3].paragraphs[0].runs[0].font.size = Pt(9)
        meta_table.columns[1].width = Cm(8)
        meta_table.columns[3].width = Cm(3)

        # Merge description cells
        meta_table.rows[0].cells[1].merge(meta_table.rows[0].cells[1])

        # Test steps
        p = doc.add_paragraph()
        run = p.add_run('Test Steps:')
        run.bold = True
        run.font.size = Pt(9)
        for line in steps.split('\n'):
            doc.add_paragraph(line.strip(), style='List Number')

        # Expected result
        p = doc.add_paragraph()
        run = p.add_run('Expected Result:')
        run.bold = True
        run.font.size = Pt(9)
        for line in expected.split('\n'):
            doc.add_paragraph(line.strip(), style='List Bullet')

        # Notes
        if notes:
            p = doc.add_paragraph()
            run = p.add_run('Notes: ')
            run.bold = True
            run.font.size = Pt(9)
            run2 = p.add_run(notes)
            run2.font.size = Pt(9)

        doc.add_paragraph()  # spacing between test cases

    doc.save(f'{OUTPUT_DIR}/02-UAT-Test-Cases.docx')
    print(f"✓ Generated: 02-UAT-Test-Cases.docx")


# ============================================================
# DOCUMENT 3: UAT SIGN-OFF
# ============================================================

def generate_signoff():
    doc = Document()
    style_document(doc)

    title = doc.add_heading('User Acceptance Testing — Sign-off', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('InvoiceDirect Mass Mailer System')
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(f'\nVersion: 2.0\nDate: {datetime.date.today().strftime("%d %B %Y")}').font.size = Pt(11)

    doc.add_page_break()

    # Test Summary
    add_heading(doc, '1. Test Summary', level=1)
    doc.add_paragraph(
        'User Acceptance Testing (UAT) has been completed for the InvoiceDirect Mass Mailer system. '
        'The following sections capture the execution summary, defect status, and formal approval.'
    )

    # Results Summary Table
    add_heading(doc, '2. Execution Summary', level=1)
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_header_row(table, ['Feature', 'Total Cases', 'Pass', 'Fail', 'Blocked'])

    features = [
        ('F-01: Customer Capture (single)', 2, '', '', ''),
        ('F-02: Customer Bulk Import (CSV)', 3, '', '', ''),
        ('F-03: Fiscalised Invoices in Folder', 4, '', '', ''),
        ('F-04: Pick PDF & Send Email', 3, '', '', ''),
        ('F-05: Customer Must Be On Platform', 2, '', '', ''),
        ('F-06: Custom Email Body Template', 2, '', '', ''),
        ('F-07: Auto-CC Sibling Contacts', 1, '', '', ''),
        ('F-08: View Sent Invoice PDF', 1, '', '', ''),
        ('F-09: Email Delivery Monitoring', 3, '', '', ''),
    ]

    for feat, total, pass_, fail, blocked in features:
        add_data_row(table, [feat, total, pass_, fail, blocked])
        set_cell_shading(table.rows[-1].cells[0], "F5F5F5")

    # Totals row
    total_row = table.add_row()
    total_row.cells[0].text = 'TOTAL'
    for i in [1, 2, 3, 4]:
        total_row.cells[i].text = ''
    for cell in total_row.cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
        set_cell_shading(cell, "1F4E79")
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    doc.add_paragraph()

    # Pass Rate
    add_heading(doc, '3. Pass Rate', level=1)
    doc.add_paragraph(
        'Pass Rate: _____ % (Passed / Total Executed × 100)'
    )
    doc.add_paragraph(
        'Minimum required pass rate: 95%'
    )

    # Defect Summary
    add_heading(doc, '4. Defect Summary', level=1)
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_header_row(table, ['Severity', 'Open', 'Resolved', 'Total'])
    severities = [
        ('Critical (P0)', '', '', ''),
        ('High (P1)', '', '', ''),
        ('Medium (P2)', '', '', ''),
        ('Low (P3)', '', '', ''),
    ]
    for sev, open_, resolved, total in severities:
        add_data_row(table, [sev, open_, resolved, total])
    # Total row
    total_row = table.add_row()
    total_row.cells[0].text = 'TOTAL'
    for i in [1, 2, 3]:
        total_row.cells[i].text = ''
    for cell in total_row.cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
        set_cell_shading(cell, "1F4E79")
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    doc.add_paragraph()

    # Outstanding Issues
    add_heading(doc, '5. Outstanding Issues / Exceptions', level=1)
    doc.add_paragraph('(List any known issues not resolved before sign-off, with agreed workarounds or acceptance criteria)')
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_header_row(table, ['Issue Ref', 'Description', 'Workaround', 'Accepted By'])
    add_data_row(table, ['', '', '', ''])

    doc.add_paragraph()

    # Sign-off
    add_heading(doc, '6. Sign-off', level=1)
    doc.add_paragraph(
        'By signing below, the stakeholders confirm that the InvoiceDirect Mass Mailer system has been reviewed '
        'and meets the specified business requirements. The system is approved for deployment to production.'
    )

    # Conditions
    add_heading(doc, 'Conditions of Acceptance', level=2)
    conditions = [
        'All Critical (P0) and High (P1) defects are resolved or explicitly waived.',
        'The pass rate meets or exceeds the 95% threshold.',
        'Any outstanding issues are documented with agreed workarounds.',
    ]
    for c in conditions:
        doc.add_paragraph(c, style='List Bullet')

    doc.add_paragraph()

    # Signature table
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_header_row(table, ['Role', 'Name', 'Signature', 'Date'])
    signatories = [
        ('Business Stakeholder', '', '', ''),
        ('Technical Lead', '', '', ''),
        ('QA / Test Coordinator', '', '', ''),
    ]
    for role, name, sig, date in signatories:
        add_data_row(table, [role, name, sig, date])

    doc.add_paragraph()
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('Decision:  [ ] Approved  [ ] Approved with Conditions  [ ] Rejected')
    run.bold = True

    doc.save(f'{OUTPUT_DIR}/03-UAT-Signoff.docx')
    print(f"✓ Generated: 03-UAT-Signoff.docx")


# ============================================================
# MAIN
# ============================================================

if __name__ == '__main__':
    import os
    os.makedirs(OUTPUT_DIR, exist_ok=True)

    print("\n=== InvoiceDirect Mass Mailer — UAT Document Generation ===\n")

    generate_test_plan()
    generate_test_cases()
    generate_signoff()

    print(f"\n✓ All documents saved to: {OUTPUT_DIR}/\n")
