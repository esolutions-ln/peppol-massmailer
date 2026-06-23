#!/usr/bin/env python3
"""
Convert Markdown developer guide to DOCX format with proper formatting.
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import re

def parse_markdown_to_docx(md_file, docx_file):
    """Convert Markdown to DOCX with formatting."""
    
    # Create document
    doc = Document()
    
    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Read markdown content
    with open(md_file, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    
    # Process line by line
    in_code_block = False
    code_lang = None
    code_lines = []
    in_table = False
    table_data = []
    
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        
        # Skip empty lines at start
        if not line and not doc.paragraphs:
            i += 1
            continue

        
        # Code blocks
        if line.startswith('```'):
            if not in_code_block:
                # Start code block
                in_code_block = True
                code_lang = line[3:].strip()
                code_lines = []
            else:
                # End code block
                in_code_block = False
                if code_lines:
                    p = doc.add_paragraph()
                    p.style = 'List Number' if doc.styles else 'Normal'
                    run = p.add_run('\n'.join(code_lines))
                    run.font.name = 'Courier New'
                    run.font.size = Pt(9)
                    p.paragraph_format.left_indent = Inches(0.5)
                    p.paragraph_format.space_before = Pt(6)
                    p.paragraph_format.space_after = Pt(6)
            i += 1
            continue
        
        if in_code_block:
            code_lines.append(line)
            i += 1
            continue
        
        # Headings
        if line.startswith('#'):
            level = len(line) - len(line.lstrip('#'))
            text = line.lstrip('#').strip()
            
            p = doc.add_heading(text, level=min(level, 9))
            if level == 1:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.size = Pt(24)
                    run.font.bold = True
            elif level == 2:
                for run in p.runs:
                    run.font.size = Pt(18)
                    run.font.color.rgb = RGBColor(0, 70, 127)
            i += 1
            continue
        
        # Horizontal rule
        if line.strip() in ['---', '___', '***']:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(12)
            i += 1
            continue
        
        # Tables
        if '|' in line and line.strip().startswith('|'):
            if not in_table:
                in_table = True
                table_data = []
            
            # Parse table row
            cells = [cell.strip() for cell in line.split('|')[1:-1]]
            
            # Skip separator row
            if all(re.match(r'^:?-+:?$', cell) for cell in cells):
                i += 1
                continue
            
            table_data.append(cells)
            
            # Check if next line is still table
            if i + 1 < len(lines) and '|' not in lines[i + 1]:
                # End of table - create it
                if table_data and len(table_data) > 1:
                    table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
                    table.style = 'Light Grid Accent 1'
                    
                    for row_idx, row_data in enumerate(table_data):
                        for col_idx, cell_text in enumerate(row_data):
                            cell = table.rows[row_idx].cells[col_idx]
                            cell.text = cell_text
                            # Bold header row
                            if row_idx == 0:
                                for paragraph in cell.paragraphs:
                                    for run in paragraph.runs:
                                        run.font.bold = True
                
                in_table = False
                table_data = []
            
            i += 1
            continue

        
        # Bullet lists
        if line.strip().startswith('- ') or line.strip().startswith('* '):
            text = line.strip()[2:]
            p = doc.add_paragraph(text, style='List Bullet')
            i += 1
            continue
        
        # Numbered lists
        match = re.match(r'^(\d+)\.\s+(.+)$', line.strip())
        if match:
            text = match.group(2)
            p = doc.add_paragraph(text, style='List Number')
            i += 1
            continue
        
        # Regular paragraph with inline formatting
        if line.strip():
            p = doc.add_paragraph()
            
            # Process inline formatting
            text = line.strip()
            
            # Simple bold/italic/code handling
            parts = re.split(r'(\*\*.*?\*\*|\*.*?\*|`.*?`|\[.*?\]\(.*?\))', text)
            
            for part in parts:
                if not part:
                    continue
                
                run = None
                
                # Bold
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.font.bold = True
                # Italic
                elif part.startswith('*') and part.endswith('*'):
                    run = p.add_run(part[1:-1])
                    run.font.italic = True
                # Code
                elif part.startswith('`') and part.endswith('`'):
                    run = p.add_run(part[1:-1])
                    run.font.name = 'Courier New'
                    run.font.size = Pt(10)
                # Link
                elif part.startswith('[') and ')' in part:
                    link_match = re.match(r'\[(.*?)\]\((.*?)\)', part)
                    if link_match:
                        link_text = link_match.group(1)
                        run = p.add_run(link_text)
                        run.font.color.rgb = RGBColor(0, 0, 255)
                        run.font.underline = True
                # Plain text
                else:
                    run = p.add_run(part)
        
        i += 1
    
    # Save document
    doc.save(docx_file)
    print(f"✅ Successfully converted to {docx_file}")
    print(f"   Total paragraphs: {len(doc.paragraphs)}")
    print(f"   Total tables: {len(doc.tables)}")

if __name__ == '__main__':
    parse_markdown_to_docx('PEPPOL_Developer_Guide.md', 'InvoiceDirect_PEPPOL_Developer_Guide.docx')
